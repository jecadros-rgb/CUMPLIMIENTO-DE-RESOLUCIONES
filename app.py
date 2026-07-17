from __future__ import annotations

import io
import base64
import hashlib
import json
import os
import re
import shutil
import tempfile
import unicodedata
import zipfile
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Any

import pandas as pd
import streamlit as st
import streamlit.components.v1 as components
from dotenv import load_dotenv
from google import genai
from google.genai import types

BASE = Path(__file__).resolve().parent
APP_VERSION = "2026.07.14-35"
FUENTES = BASE / "fuentes_permanentes"
INSTRUCCIONES = BASE / "instrucciones" / "instrucciones_juridicas.txt"
CRITERIOS_INSTRUCCION = BASE / "instrucciones" / "criterios_evaluacion_obligatorios.txt"
TEMP = BASE / "expedientes_temporales"
SALIDAS = BASE / "salidas"
HISTORIAL = SALIDAS / "evaluaciones.xlsx"
FUENTES_REQUERIDAS = [
    "PLANTILLAS cumplimiento.docx",
    "PLANTILLAS DENUNCIAS ACTUALIZADAS.docx",
    "notificaciones 2026.csv",
    "CONTADOR DE PLAZOS - TRASU 2026.xlsx",
    "PAUTAS PAS.xlsx",
]
COLUMNAS = ["Expediente", "Empresa operadora", "Tipo de acto",
    "Número de resolución o carta", "Fecha de notificación o emisión",
    "Fecha máxima de vencimiento", "Obligación principal", "Estado de ejecución",
    "Fecha de ejecución acreditada", "Resultado",
    "Tipo de incumplimiento", "Subsanación voluntaria", "PAS / NO PAS", "Sustento breve",
    "Párrafo final", "Datos pendientes", "Documentos usados", "Fecha de evaluación"]

load_dotenv(BASE / ".env")
TEMP.mkdir(exist_ok=True); SALIDAS.mkdir(exist_ok=True)

st.set_page_config(page_title="CumpleTRASU", page_icon="⚖️", layout="wide")
st.markdown("""<style>
:root{--blue:#005f9e;--blue-dark:#08466f;--cyan:#15b8d4;--sky:#eaf8fc;--orange:#f58220;--ink:#173b57;--muted:#5f7f94;--border:#bfdce8}
.stApp{background:linear-gradient(180deg,#eaf8fc 0,#ffffff 30%,#f5fbfd 100%);color:var(--ink)}
[data-testid="stHeader"]{background:rgba(255,255,255,.94);border-bottom:1px solid #d8ebf2}.block-container{padding-top:1.25rem;max-width:1600px}
.hero{border:0;border-radius:18px;padding:20px 26px;background:linear-gradient(115deg,#005f9e 0%,#0786b9 58%,#15b8d4 100%);box-shadow:0 10px 28px rgba(0,95,158,.18);margin-bottom:18px}
.brand{font-size:2rem;font-weight:800;color:#fff}.brand b{color:#ffb15d}.sub{color:#e8f8fc}
.tag{float:right;border:1px solid rgba(255,255,255,.65);background:rgba(255,255,255,.13);border-radius:999px;padding:7px 12px;color:#fff}
.panel{background:#fff;border:1px solid var(--border);border-radius:14px;padding:14px;margin-bottom:12px;box-shadow:0 5px 16px rgba(0,95,158,.06)}
.light{color:var(--blue);font-size:.82rem;font-weight:750;text-transform:uppercase;letter-spacing:.08em}
.traffic{font-size:1.25rem;font-weight:750;padding:12px;border-radius:10px;text-align:center}
.ok{background:#d9f7fb;color:#006b7b;border:1px solid #8bdce8}.bad{background:#fff0e3;color:#b55300;border:1px solid #ffc285}.wait{background:#e7f2fb;color:#075f96;border:1px solid #a9d3ee}
[data-testid="stWidgetLabel"] p,[data-testid="stMarkdownContainer"] p,[data-testid="stCaptionContainer"],h1,h2,h3{color:var(--ink)}
[data-testid="stCaptionContainer"]{color:var(--muted)!important}
[data-testid="stTextInput"] input,[data-testid="stTextArea"] textarea{background:#fff!important;color:var(--ink)!important;border:1px solid var(--border)!important;border-radius:10px!important;box-shadow:0 2px 8px rgba(0,95,158,.05)}
[data-testid="stSelectbox"] div[data-baseweb="select"]>div{background:#fff!important;color:var(--ink)!important;border-color:var(--border)!important;border-radius:10px!important}
[data-testid="stSelectbox"] svg{fill:var(--blue)!important}
div[data-testid="stFileUploader"]{background:#fff;border:1px dashed var(--cyan);border-radius:14px;padding:8px;box-shadow:0 4px 14px rgba(0,95,158,.06)}
[data-testid="stFileUploaderDropzone"]{background:#f4fbfe!important}
[data-testid="stExpander"]{background:#fff;border:1px solid var(--border);border-radius:12px;box-shadow:0 3px 12px rgba(0,95,158,.05)}
[data-testid="stMetric"]{background:#fff;border:1px solid var(--border);border-radius:12px;padding:10px}
hr{border-color:#d6eaf2!important}
.stButton>button,.stDownloadButton>button{border-radius:10px;border:1px solid var(--blue);background:#fff;color:var(--blue);font-weight:700}.stDownloadButton>button{width:100%}
.stButton>button:hover,.stDownloadButton>button:hover{border-color:var(--cyan);color:var(--blue-dark);background:#effaff}
.stButton>button[kind="primary"]{background:linear-gradient(135deg,var(--orange),#ff9f43);color:#fff;border:0;box-shadow:0 6px 15px rgba(245,130,32,.25)}
.stButton>button[kind="primary"]:hover{background:linear-gradient(135deg,#e66f0e,var(--orange));color:#fff}
/* La ayuda interna de objetos (DeltaGenerator) no forma parte de CumpleTRASU. */
[data-testid="stHelp"]{display:none!important}
</style>""", unsafe_allow_html=True)
st.markdown('<div class="hero"><span class="tag">OSIPTEL / TRASU</span><div class="brand">Cumple<b>TRASU</b></div><div class="sub">Evaluador Automatizado de Cumplimiento de Resoluciones</div></div>', unsafe_allow_html=True)
st.caption(f"Versión {APP_VERSION}")

def get_api_key() -> str | None:
    try: return st.secrets.get("GEMINI_API_KEY") or os.getenv("GEMINI_API_KEY")
    except Exception: return os.getenv("GEMINI_API_KEY")

def gemini_client():
    return genai.Client(api_key=get_api_key())

def gemini_text(system: str, user: Any, json_mode: bool=False) -> str:
    config=types.GenerateContentConfig(system_instruction=system,max_output_tokens=32768)
    if json_mode: config.response_mime_type="application/json"
    client=gemini_client()
    errors=[]
    # Flash-Lite first: it is both far cheaper and has a much higher free-tier
    # daily quota than 3.5-flash, so prefer it in every call, not only OCR.
    models=("gemini-3.1-flash-lite","gemini-3.5-flash")
    for model in models:
        # Do not repeat an identical rejected request: repeated RECITATION calls
        # consume the small free quota without improving the document extraction.
        for attempt in range(1):
            try:
                response=client.models.generate_content(model=model,contents=user,config=config)
                candidates=response.candidates or []
                finish=str(candidates[0].finish_reason) if candidates else "SIN_CANDIDATOS"
                text=response.text if candidates and candidates[0].content and candidates[0].content.parts else None
                if not text: raise RuntimeError(f"El modelo {model} no devolvió contenido utilizable (motivo: {finish}).")
                return text
            except Exception as e:
                errors.append(f"{model} (intento {attempt+1}): {e}")
                if "RECITATION" not in str(e): break
    raise RuntimeError(" | ".join(errors))

def parse_json_response(text: str) -> Any:
    """Gemini's JSON mode can still wrap output in markdown fences or append trailing text; recover the first valid value."""
    cleaned=text.strip()
    if cleaned.startswith("```"):
        cleaned=re.sub(r"^```[a-zA-Z]*\n?","",cleaned)
        cleaned=re.sub(r"```\s*$","",cleaned).strip()
    return json.JSONDecoder().raw_decode(cleaned)[0]

def json_object(value: Any, preferred_keys: tuple[str,...]=()) -> dict[str,Any]:
    """Recover a JSON object from harmless model wrappers without another API request."""
    current=value
    for _ in range(4):
        if isinstance(current,dict):
            for key in preferred_keys:
                nested=current.get(key)
                if isinstance(nested,dict):
                    current=nested
                    break
            else:
                return current
            continue
        if isinstance(current,list):
            dictionaries=[item for item in current if isinstance(item,dict)]
            if not dictionaries: return {}
            # Gemini occasionally wraps the requested object in a one-element
            # array. Prefer the object containing the expected top-level keys.
            expected=set(preferred_keys)|{"acto","obligaciones","medios_probatorios","ficha",
                                          "evaluacion_juridica","resultado","parrafo_final"}
            current=max(dictionaries,key=lambda item:len(expected & set(item)))
            continue
        if isinstance(current,str):
            try: current=parse_json_response(current)
            except Exception: return {}
            continue
        return {}
    return current if isinstance(current,dict) else {}

def fuente_status() -> tuple[bool, list[str]]:
    faltan = [n for n in FUENTES_REQUERIDAS if not (FUENTES / n).is_file()]
    instrucciones_ok = INSTRUCCIONES.is_file() and INSTRUCCIONES.stat().st_size > 150
    if instrucciones_ok:
        instrucciones_ok = "PEGUE AQUÍ" not in INSTRUCCIONES.read_text("utf-8", errors="ignore")
    if not instrucciones_ok: faltan.append("instrucciones_juridicas.txt (contenido real)")
    if not CRITERIOS_INSTRUCCION.is_file() or CRITERIOS_INSTRUCCION.stat().st_size < 500:
        faltan.append("criterios_evaluacion_obligatorios.txt")
    return not faltan, faltan

def safe_name(name: str) -> str:
    return re.sub(r"[^\w.() -]", "_", Path(name).name, flags=re.UNICODE)[:180]

def extract_archives(folder: Path) -> list[Path]:
    out=[]
    for f in list(folder.iterdir()):
        try:
            target=folder/(f.stem+"_extraido"); target.mkdir(exist_ok=True)
            if f.suffix.lower()==".zip":
                with zipfile.ZipFile(f) as z:
                    for m in z.infolist():
                        dest=(target/m.filename).resolve()
                        if target.resolve() not in dest.parents and dest != target.resolve(): continue
                        z.extract(m,target)
            elif f.suffix.lower()==".7z":
                import py7zr
                with py7zr.SevenZipFile(f) as z: z.extractall(target)
            elif f.suffix.lower()==".rar":
                import rarfile
                with rarfile.RarFile(f) as z: z.extractall(target)
            else: continue
            out += [p for p in target.rglob("*") if p.is_file()]
        except Exception as e: st.warning(f"No se pudo descomprimir {f.name}: {e}")
    return out

def read_file(path: Path) -> str:
    ext=path.suffix.lower()
    try:
        if ext==".pdf":
            from pypdf import PdfReader
            reader=PdfReader(str(path))
            text="\n".join((p.extract_text() or "") for p in reader.pages)
            # Hybrid PDFs often contain a short text layer plus printers,
            # screenshots or scanned evidence as images. A text-length check
            # alone silently omitted those facts. OCR every page locally when
            # raster images are present and append it to the extracted text.
            try:
                has_raster_images=any(bool(getattr(page,"images",[])) for page in reader.pages)
            except Exception:
                has_raster_images=False
            if len(text.strip())<80 or has_raster_images:
                try:
                    import pytesseract
                    from pdf2image import convert_from_path
                    images=convert_from_path(str(path),dpi=170,thread_count=2)
                    ocr_text="\n".join(
                        f"--- OCR PÁGINA {i+1} ---\n"+pytesseract.image_to_string(image,lang="spa")
                        for i,image in enumerate(images))
                    if ocr_text.strip(): text=(text+"\n\n"+ocr_text).strip()
                except Exception: pass
            if len(text.strip())<80:
                return vision_ocr(path)
            return text
        if ext==".docx":
            from docx import Document
            d=Document(path); return "\n".join([p.text for p in d.paragraphs]+[" | ".join(c.text for c in r.cells) for t in d.tables for r in t.rows])
        if ext in {".xlsx",".xls"}:
            book=pd.read_excel(path,sheet_name=None,dtype=str)
            return "\n".join(f"[{k}]\n{v.fillna('').to_csv(index=False)}" for k,v in book.items())
        if ext==".csv": return pd.read_csv(path,dtype=str).fillna("").to_csv(index=False)
        if ext in {".txt",".md"}: return path.read_text("utf-8",errors="ignore")
        if ext in {".png",".jpg",".jpeg",".tif",".tiff",".bmp"}:
            try:
                import pytesseract
                from PIL import Image, ImageSequence
                source=Image.open(path)
                text="\n".join(pytesseract.image_to_string(frame.convert("RGB"),lang="spa")
                               for frame in ImageSequence.Iterator(source))
                legal_markers=("resuelve","declarar fundad","declara fundad","artículo 1","articulo 1","sentido de la resolución")
                if text.strip() and any(x in text.lower() for x in legal_markers): return text
            except Exception: pass
            return vision_ocr(path)
    except Exception as e: return f"[Error al leer {path.name}: {e}]"
    return ""

def vision_ocr(path: Path) -> str:
    """OCR scanned images in small batches, prioritizing the end of legal resolutions."""
    schema={"tipo_documento":"","expediente":"","numero_acto":"","fecha_acto":"",
            "ha_resuelto":[{"numeral":"","obligacion_o_disposicion":"","plazo":"",
                             "destinatario":"","es_obligacion_accesoria":False}],
            "hechos_probatorios_relevantes":[]}
    instruction="""Analiza las imágenes de este documento jurídico y extrae hechos estructurados; NO transcribas íntegramente el documento ni reproduzcas párrafos largos.
Identifica tipo de documento, expediente, número y fecha del acto. Si aparece HA RESUELTO, extrae por separado cada numeral, describiendo fielmente pero de forma concisa su obligación o disposición, su plazo, destinatario y si es una obligación accesoria de informar posteriormente el cumplimiento al usuario o al TRASU. No inventes datos ilegibles.
Para documentos probatorios, resume únicamente fechas, importes, acciones efectivamente ejecutadas y constancias visibles. Devuelve JSON conforme al esquema."""
    def facts_text(raw: str) -> str:
        data=parse_json_response(raw) if raw else {}
        if not isinstance(data,dict): return ""
        lines=[f"TIPO DE DOCUMENTO: {data.get('tipo_documento','')}",
               f"EXPEDIENTE: {data.get('expediente','')}",f"NÚMERO DE ACTO: {data.get('numero_acto','')}",
               f"FECHA DEL ACTO: {data.get('fecha_acto','')}"]
        operative=data.get("ha_resuelto") or []
        if operative:
            lines.append("HA RESUELTO")
            for item in operative:
                if not isinstance(item,dict): continue
                lines.append(f"NUMERAL {item.get('numeral','')}: {item.get('obligacion_o_disposicion','')} | PLAZO: {item.get('plazo','')} | DESTINATARIO: {item.get('destinatario','')} | ACCESORIA: {item.get('es_obligacion_accesoria',False)}")
        for fact in (data.get("hechos_probatorios_relevantes") or []): lines.append(f"HECHO PROBATORIO: {fact}")
        return "\n".join(lines)
    try:
        from PIL import Image
        Image.MAX_IMAGE_PIXELS=None
        source=Image.open(path)
        total=max(1,int(getattr(source,"n_frames",1)))
        # HA RESUELTO is normally on the final pages. For very large TIFF files,
        # read the cover and the final 18 pages instead of silently stopping at
        # the first 20 pages, which previously omitted the operative section.
        if total<=30:
            indices=list(range(total))
        else:
            indices=sorted(set(list(range(min(6,total)))+list(range(max(0,total-18),total))))
        batches=[indices[i:i+5] for i in range(0,len(indices),5)]
        transcriptions=[]; errors=[]
        for batch in batches:
            content=[instruction+"\nESQUEMA: "+json.dumps(schema,ensure_ascii=False)]
            for index in batch:
                source.seek(index)
                image=source.convert("RGB"); image.thumbnail((1500,1500))
                buf=io.BytesIO(); image.save(buf,format="JPEG",quality=78,optimize=True)
                content.append(f"--- Página {index+1} de {total} ---")
                content.append(types.Part.from_bytes(data=buf.getvalue(),mime_type="image/jpeg"))
            try:
                text=gemini_text("Eres un extractor de hechos jurídicos visibles en imágenes. Evita transcripciones extensas.",content,json_mode=True)
                converted=facts_text(text)
                if converted.strip(): transcriptions.append(converted)
            except Exception as batch_error:
                errors.append(f"páginas {batch[0]+1}-{batch[-1]+1}: {batch_error}")
        combined_transcription="\n\n".join(transcriptions)
        # Recovery used only when the ordinary OCR read a TRASU resolution but
        # omitted its operative section.  Inspect the final pages individually,
        # in reverse order, and stop as soon as HA RESUELTO is recovered.  This
        # supplements the first extraction; it does not replace any other data.
        looks_like_trasu=("trasu" in path.name.lower() or
                          "resoluci" in path.name.lower() or
                          "TRASU" in combined_transcription.upper())
        if looks_like_trasu and "HA RESUELTO" not in combined_transcription.upper():
            operative_schema={"ha_resuelto":[{
                "numeral":"", "obligacion_o_disposicion":"", "plazo":"",
                "destinatario":"", "es_obligacion_accesoria":False
            }]}
            operative_instruction="""Examina solamente esta pagina de una resolucion TRASU.
Si contiene el titulo HA RESUELTO, SE RESUELVE o la continuacion de sus numerales, resume TODOS los numerales visibles en ha_resuelto. Conserva la relacion entre el mandato material y su plazo aunque aparezcan en numerales consecutivos. Incluye tambien las obligaciones accesorias y marcalas como tales. No transcribas parrafos completos y no inventes texto. Si esta pagina no contiene parte resolutiva, devuelve ha_resuelto como lista vacia. Devuelve solo JSON conforme al esquema."""
            recovery_errors=[]
            # One recovery request at most. Send the central body of the final
            # pages together instead of making one or two API calls per page.
            # This keeps the rule general while protecting the user's quota.
            try:
                recovery_content=[operative_instruction+"\nESQUEMA: "+json.dumps(operative_schema,ensure_ascii=False)]
                for index in indices[-6:]:
                    source.seek(index)
                    original=source.convert("RGB")
                    width,height=original.size
                    body=original.crop((0,int(height*.12),width,int(height*.92))) if height>800 else original
                    body.thumbnail((1500,1500))
                    buf=io.BytesIO(); body.save(buf,format="JPEG",quality=80,optimize=True)
                    recovery_content.append(f"Pagina {index+1} de {total}, cuerpo central")
                    recovery_content.append(types.Part.from_bytes(data=buf.getvalue(),mime_type="image/jpeg"))
                # The previous working version recovered this section as plain
                # text. JSON extraction may legally succeed while returning an
                # empty ha_resuelto list, so use one concise plain-text request
                # here and normalize its numbered output under the heading.
                recovery_content[0]="""Localiza en estas paginas la parte resolutiva de la Resolucion TRASU.
Devuelve exclusivamente un resumen fiel con este formato:
HA RESUELTO
NUMERAL 1: disposicion y plazo visible
NUMERAL 2: disposicion y plazo visible
Incluye todos los numerales resolutivos visibles y conserva las referencias entre numerales consecutivos. No agregues antecedentes, considerandos ni hechos de otros documentos. Si no aparece ninguna parte resolutiva, responde exactamente NO_ENCONTRADO."""
                raw=gemini_text("Lee visualmente la parte resolutiva y resume sus numerales sin copiar parrafos extensos.",recovery_content,json_mode=False)
                normalized_raw=str(raw or "").strip()
                if normalized_raw and "NO_ENCONTRADO" not in normalized_raw.upper() and re.search(r"\bNUMERAL\s+\w+",normalized_raw,re.I):
                    if "HA RESUELTO" not in normalized_raw.upper():
                        normalized_raw="HA RESUELTO\n"+normalized_raw
                    combined_transcription=(combined_transcription+"\n\n"+normalized_raw).strip()
            except Exception as recovery_error:
                recovery_errors.append(f"recuperacion conjunta: {recovery_error}")
        if combined_transcription:
            return combined_transcription
        # Second route: let the Files API ingest the original TIFF. This avoids
        # Pillow/libtiff decoder limitations that affect some official scans.
        try:
            client=gemini_client()
            uploaded=client.files.upload(file=str(path))
            prompt=instruction+"\nESQUEMA: "+json.dumps(schema,ensure_ascii=False)
            raw=gemini_text("Eres un extractor de hechos jurídicos. Evita transcripciones extensas.",[uploaded,prompt],json_mode=True)
            converted=facts_text(raw)
            if converted.strip(): return converted
            errors.append("el archivo TIFF original no produjo texto")
        except Exception as upload_error:
            errors.append(f"lectura directa del TIFF: {upload_error}")
        raise RuntimeError("; ".join(errors) if errors else "ningún método produjo texto")
    except Exception as e:
        # If page decoding itself failed, still try the original file once.
        try:
            client=gemini_client()
            uploaded=client.files.upload(file=str(path))
            prompt=instruction+"\nESQUEMA: "+json.dumps(schema,ensure_ascii=False)
            raw=gemini_text("Eres un extractor de hechos jurídicos. Evita transcripciones extensas.",[uploaded,prompt],json_mode=True)
            converted=facts_text(raw)
            if converted.strip(): return converted
        except Exception as upload_error:
            return f"[OCR no disponible para {path.name}: conversión por páginas: {e}; lectura directa: {upload_error}]"
        return f"[OCR no disponible para {path.name}: {e}]"

def classify(name: str, text: str) -> str:
    # TRASU must win when it is present in the filename or document text.
    # SAR/SARA/SAP are acronyms: substring matching wrongly classified ordinary
    # words containing "sar" and caused valid TRASU resolutions to be skipped.
    s=name+" "+text[:6000]
    folded=unicodedata.normalize("NFD",s.upper())
    folded="".join(c for c in folded if unicodedata.category(c)!="Mn")
    if re.search(r"(?<![A-Z0-9])TRASU(?![A-Z0-9])",folded) or "TRASU" in name.upper():
        return "Resolución TRASU"
    for acronym,label in (("SARA","SARA"),("SAR","SAR"),("SAP","SAP")):
        if re.search(rf"(?<![A-Z0-9]){acronym}(?![A-Z0-9])",folded): return label
    low=folded.lower()
    for key,label in (("denuncia","Denuncia"),("carta","Carta"),("primera instancia","Resolución de primera instancia")):
        if key in low: return label
    return "No identificado"

def extract_resolutive_part(documents: dict[str,str]) -> str | None:
    """Extract the operative section; obligations may not be inferred elsewhere."""
    candidates=[]
    for name,text in documents.items():
        if classify(name,text)!="Resolución TRASU": continue
        upper=unicodedata.normalize("NFD",text.upper())
        upper="".join(c for c in upper if unicodedata.category(c)!="Mn")
        upper_name=unicodedata.normalize("NFD",name.upper())
        upper_name="".join(c for c in upper_name if unicodedata.category(c)!="Mn")
        resolution_document=(
            bool(re.search(r"RESOLUCION[^\n]{0,40}FINAL",upper_name)) or
            bool(re.search(r"TRIBUNAL\s+ADMINISTRATIVO[\s\S]{0,500}RESOLUCION\s+FINAL",upper))
        )
        # A letter or an evidence annex may mention TRASU, but that does not
        # turn it into the resolution whose operative section defines the duty.
        if not resolution_document: continue
        anchors=[]
        for pattern in (r"\bHA\s+RESUELTO\b",r"\bSE\s+RESUELVE\b",r"\bRESUELVE\s*:",r"\bPARTE\s+RESOLUTIVA\b",
                        r"\bARTICULO\s+(?:PRIMERO|1)\b",r"\bART\.\s*1\b"):
            match=re.search(pattern,upper)
            if match: anchors.append(match.start())
        if anchors:
            start=min(anchors); section=text[start:start+18000]
            if re.search(r"DECLARAR\s+(?:EL\s+RECLAMO\s+)?FUNDAD",section,re.I):
                candidates.append(f"### {name} — PARTE RESOLUTIVA\n{section}")
            else:
                candidates.append(f"### {name} — PARTE RESOLUTIVA\n{section}")
            continue
        # OCR sometimes omits the heading but preserves the operative declaration.
        match=re.search(r"(?:DECLARAR|DECLARA|SE\s+DECLARA)[^.\n]{0,180}?FUNDAD[OA]",upper)
        if match: candidates.append(f"### {name} — DECLARACIÓN FUNDADA\n{text[max(0,match.start()-800):match.start()+12000]}")
        elif re.search(r"SENTIDO\s+DE\s+LA\s+RESOLUCION[^.\n]{0,100}FUNDAD[OA]",upper):
            # Some official copies place the operative articles on the final pages.
            candidates.append(f"### {name} — PÁGINAS FINALES DE RESOLUCIÓN FUNDADA\n{text[-18000:]}")
        elif len(text.strip())>500:
            # Last-resort candidate is restricted to the end of the resolution,
            # never to letters, briefs or evidence from other documents.
            candidates.append(f"### {name} — CANDIDATO DE PARTE RESOLUTIVA (PÁGINAS FINALES)\n{text[-18000:]}")
    return "\n\n".join(candidates) if candidates else None

def extract_trasu_mandate_and_term(resolutive: str) -> dict[str,Any]:
    """Read HA RESUELTO as a whole and associate the principal mandate with its own term."""
    schema={
        "obligacion_principal":"",
        "plazo_principal_textual":"",
        "numero_plazo":None,
        "tipo_plazo":"dias_habiles|dias_calendario|meses|no_identificado",
        "numerales_fuente":[],
        "obligaciones_accesorias_excluidas":[],
    }
    system="""Lee íntegramente la sección HA RESUELTO de una Resolución TRASU.
Identifica la obligación material principal relacionada con la prestación o controversia del servicio y el plazo expresamente asociado a SU ejecución. Lee conjuntamente todos los numerales: la obligación y su plazo pueden estar en puntos consecutivos.
No uses una lista cerrada de verbos y no exijas que el mandato esté en el mismo numeral que declara fundado el reclamo.
Excluye las obligaciones posteriores de informar el cumplimiento al usuario, informar al TRASU, remitir constancias, acreditar comunicaciones o reportar acciones, junto con los plazos de esas obligaciones accesorias. Nunca confundas esos plazos con el plazo principal.
Los considerandos solo pueden aclarar una referencia del mandato resolutivo; no pueden crear ni ampliar la obligación.
Copia literalmente el plazo principal. Si no está legible, usa no_identificado; no adivines ni apliques un plazo general.
Devuelve JSON válido conforme al esquema."""
    raw=gemini_text(system,json.dumps({"esquema":schema,"ha_resuelto":resolutive},ensure_ascii=False),json_mode=True)
    data=parse_json_response(raw) if raw else {}
    if not isinstance(data,dict):
        raise ValueError("No se pudo interpretar la sección HA RESUELTO")
    if not str(data.get("obligacion_principal") or "").strip():
        raise ValueError("No se pudo identificar la obligación material principal en HA RESUELTO")
    if not str(data.get("plazo_principal_textual") or "").strip() or str(data.get("tipo_plazo") or "") == "no_identificado":
        raise ValueError("No se pudo identificar el plazo vinculado a la obligación principal en HA RESUELTO")
    return data

def exact_case_record(expediente: str) -> dict[str,str] | None:
    """Return institutional data from the exact notification row, never from a fuzzy match."""
    path=FUENTES/"notificaciones 2026.csv"
    for _,df in {"Notificaciones 2026":pd.read_csv(path,dtype=str,encoding="utf-8-sig")}.items():
        cols={str(c).strip().upper():c for c in df.columns}
        date_key="FEC_NOT_EMP_ELE_TEXTO" if "FEC_NOT_EMP_ELE_TEXTO" in cols else ("FEC_NOT_EMP_ELE" if "FEC_NOT_EMP_ELE" in cols else None)
        if "NRO_EXPEDIENTE" in cols and date_key:
            values=df[cols["NRO_EXPEDIENTE"]].fillna("").astype(str).str.strip()
            target=re.sub(r"[^A-Z0-9]","",expediente.upper())
            normalized=values.str.upper().str.replace(r"[^A-Z0-9]","",regex=True)
            hit=df.loc[normalized==target]
            if not hit.empty and pd.notna(hit.iloc[0][cols[date_key]]):
                row=hit.iloc[0]
                raw=str(row[cols[date_key]]).strip()
                # The protected annual source stores Peruvian dates as dd/mm/yyyy.
                parsed=pd.to_datetime(raw,dayfirst=True,errors="coerce")
                notification=parsed.strftime("%d/%m/%Y") if pd.notna(parsed) else raw
                operator=""
                if "EMPRESA" in cols and pd.notna(row[cols["EMPRESA"]]):
                    operator=re.sub(r"\s+"," ",str(row[cols["EMPRESA"]]).strip())
                    if operator.lower() in {"nan","none","no identificado"}: operator=""
                return {"fecha_notificacion":notification,"empresa_operadora":operator}
    return None

def identify_exact_expediente(context: str) -> str | None:
    """Resolve filename-safe variants to one exact value stored in NRO_EXPEDIENTE."""
    normalized_context=re.sub(r"[^A-Z0-9]","",context.upper())
    candidates=[]
    path=FUENTES/"notificaciones 2026.csv"
    for _,df in {"Notificaciones 2026":pd.read_csv(path,dtype=str,encoding="utf-8-sig")}.items():
        cols={str(c).strip().upper():c for c in df.columns}
        if "NRO_EXPEDIENTE" not in cols: continue
        for value in df[cols["NRO_EXPEDIENTE"]].dropna().astype(str).str.strip().unique():
            normalized=re.sub(r"[^A-Z0-9]","",value.upper())
            if len(normalized)>=10 and normalized in normalized_context:
                candidates.append(value)
    unique=list(dict.fromkeys(candidates))
    return unique[0] if len(unique)==1 else None

def parse_trasu_name(names: str) -> str | None:
    compact=re.sub(r"[^A-Z0-9]","",names.upper())
    match=re.search(r"(\d{7})(20\d{2})TRASU(STRA|STRQJ)",compact)
    if not match: return None
    suffix="ST-RA" if match.group(3)=="STRA" else "ST-RQJ"
    return f"{match.group(1)}-{match.group(2)}/TRASU/{suffix}"

def source_text(name: str, limit=50000) -> str:
    return read_file(FUENTES/name)[:limit]

def parse_excel_date(value: Any) -> pd.Timestamp:
    """Accept displayed dates, timestamps and Excel serial date numbers."""
    raw=str(value).strip()
    if re.fullmatch(r"\d+(?:\.0+)?",raw):
        serial=float(raw)
        if 20000 <= serial <= 80000:
            return pd.Timestamp("1899-12-30")+pd.to_timedelta(serial,unit="D")
    parsed=pd.to_datetime(raw,dayfirst=True,errors="coerce")
    if pd.isna(parsed): raise ValueError(f"fecha de notificación no reconocida: {raw}")
    return pd.Timestamp(parsed)

def _legal_tokens(value: Any) -> set[str]:
    text=unicodedata.normalize("NFD",str(value).lower())
    text="".join(c for c in text if unicodedata.category(c)!="Mn")
    return {x for x in re.findall(r"[a-z0-9]{4,}",text) if x not in {
        "para","como","esta","este","desde","hasta","sobre","entre","donde",
        "empresa","operadora","resolucion","trasu","usuario","cumplimiento"}}

def relevant_excel_rules(name: str, case_context: str, limit: int=28,
                         service_restriction: bool | None=None,
                         information_delivery: bool=False) -> str:
    """Retrieve applicable rows instead of truncating a whole legal workbook."""
    book=pd.read_excel(FUENTES/name,sheet_name=None,dtype=str,header=None)
    query=_legal_tokens(case_context)
    normalized_case=unicodedata.normalize("NFD",case_context.lower())
    normalized_case="".join(c for c in normalized_case if unicodedata.category(c)!="Mn")
    discount_case=any(k in normalized_case for k in (
        "descuento","ajuste","facturacion","importe","nota de credito",
        "oferta","promocion","beneficio recurrente"))
    ranked=[]
    for sheet,df in book.items():
        for idx,row in df.fillna("").iterrows():
            cells=[str(x).strip() for x in row.tolist() if str(x).strip()]
            if not cells: continue
            line=" | ".join(cells)
            tokens=_legal_tokens(line)
            score=len(query & tokens)
            low=line.lower()
            restriction_rule=("restricción del servicio" in low or
                              "restriccion del servicio" in low or
                              "materia de falta de servicio" in low)
            # La pauta que declara irreversibles los efectos de una falta de
            # servicio no es transversal. Excluirla impide aplicarla a una
            # obligación meramente informativa, contractual o económica.
            if service_restriction is False and restriction_rule:
                continue
            # Only truly transversal rules are boosted for every matter.
            # Discount/programming rules are boosted exclusively when the case
            # itself concerns billing, adjustments, offers or promotions.
            general=any(k in low for k in (
                "cese total","reversión total","subsanación voluntaria",
                "análisis por cada mandato","conclusión es por resolución",
                "cumplimiento parcial"))
            discount_rule=any(k in low for k in (
                "registro y activación de los descuentos","descuento recurrente",
                "ajuste recurrente","meses pendientes","meses restantes"))
            if information_delivery and discount_rule:
                continue
            general=general or (discount_case and discount_rule)
            if general: score+=100
            if score: ranked.append((score,sheet,int(idx)+1,line))
    ranked.sort(key=lambda x:(-x[0],x[1],x[2]))
    selected=ranked[:limit]
    return "\n".join(f"[{name} / {s} / fila {r}] {line}" for _,s,r,line in selected)

def _fold_legal_text(value: Any) -> str:
    text=unicodedata.normalize("NFD",str(value or "").lower())
    return "".join(c for c in text if unicodedata.category(c)!="Mn")

def is_information_delivery_obligation(value: Any) -> bool:
    """Identify mandates whose material performance is delivery of documents or information."""
    folded=_fold_legal_text(value)
    actions=("informar","brindar","remitir","entregar","comunicar","trasladar", "poner en conocimiento")
    objects=("contrato","terminos","condiciones","tarif","monto","informacion", "document")
    return any(x in folded for x in actions) and any(x in folded for x in objects)

def _legal_dates_in(value: Any) -> list[pd.Timestamp]:
    """Read dates from extracted evidence without depending on Gemini's prose."""
    raw=str(value or "").strip()
    if not raw:
        return []
    months={"enero":"01","febrero":"02","marzo":"03","abril":"04","mayo":"05","junio":"06",
            "julio":"07","agosto":"08","septiembre":"09","setiembre":"09","octubre":"10",
            "noviembre":"11","diciembre":"12"}
    normalized=re.sub(
        r"\b(\d{1,2})\s+de\s+("+"|".join(months)+r")\s+de\s+(\d{4})\b",
        lambda m:f"{int(m.group(1)):02d}/{months[m.group(2).lower()]}/{m.group(3)}",
        raw,flags=re.I)
    candidates=re.findall(r"\b\d{1,2}/\d{1,2}/\d{4}\b|\b\d{4}-\d{1,2}-\d{1,2}\b",normalized)
    if not candidates:
        candidates=[normalized]
    dates=[]
    for candidate in candidates:
        parsed=pd.to_datetime(candidate,dayfirst=not re.match(r"^\d{4}-",candidate),errors="coerce")
        if pd.notna(parsed):
            stamp=pd.Timestamp(parsed).normalize()
            if stamp not in dates: dates.append(stamp)
    return dates

def _criteria_rows() -> list[dict[str,Any]]:
    """Parse the institutional instruction while preserving each row's matter."""
    raw=CRITERIOS_INSTRUCCION.read_text("utf-8",errors="ignore")
    matches=list(re.finditer(r"(?ms)^FILA\s+(\d+):\s*(.*?)(?=^FILA\s+\d+:|\Z)",raw))
    known_matters={
        "Facturación y cobro","Calidad e idoneidad",
        "Incumplimiento de condiciones contractuales, ofertas y promociones",
        "Corte o Baja Injustificada","Instalación, activación o traslado del servicio",
        "Falta de ejecución de baja o traslado del servicio","Recargas",
        "Contratación no solicitada","Migración","Portabilidad","Todas las materias",
        "Suspensión por uso prohibido","Activación de promoción",
        "Devoluciones en Centros de Atención de la empresa Operadora",
        "Portabilidad no autorizada","Cambio de número no autorizado",
    }
    folded_known={_fold_legal_text(x):x for x in known_matters}
    current=""
    rows=[]
    for match in matches:
        number=int(match.group(1))
        body=re.sub(r"\s+"," ",match.group(2)).strip()
        parts=[p.strip() for p in body.split("|")]
        first=_fold_legal_text(parts[0] if parts else "")
        if first in folded_known:
            current=folded_known[first]
        matter="Notificaciones" if number==66 else current
        if number>=3:
            rows.append({"fila":number,"materia":matter,"texto":body})
    return rows

def select_applicable_criteria(obligation: str) -> tuple[str,str,list[int]]:
    """Route the verified mandate to its own matrix block before legal reasoning."""
    rows=_criteria_rows()
    folded=_fold_legal_text(obligation)
    routing={
        "Corte o Baja Injustificada":("reconect","reactiv","restablec","corte injust","baja injust","servicio suspend"),
        "Calidad e idoneidad":("averia","calidad","idoneidad","reparar servicio","operatividad del servicio"),
        "Incumplimiento de condiciones contractuales, ofertas y promociones":("oferta","promocion","beneficio","descuento","condicion contractual"),
        "Facturación y cobro":("facturacion","cobro","ajustar importe","anular importe","nota de credito","estado de cuenta"),
        "Instalación, activación o traslado del servicio":("instalar","instalacion","activar el servicio","activacion del servicio","trasladar el servicio"),
        "Falta de ejecución de baja o traslado del servicio":("dar de baja","baja del servicio","ejecutar la baja","traslado del servicio"),
        "Contratación no solicitada":("contratacion no solicitada","servicio no solicitado"),
        "Migración":("migracion","migrar","cambio de plan"),
        "Portabilidad no autorizada":("portabilidad no autorizada",),
        "Portabilidad":("portabilidad","portar el numero"),
        "Cambio de número no autorizado":("cambio de numero no autorizado",),
        "Recargas":("recarga","saldo recargado"),
        "Suspensión por uso prohibido":("suspension por uso prohibido","rentseg","imei"),
        "Activación de promoción":("activar promocion","activacion de promocion"),
        "Devoluciones en Centros de Atención de la empresa Operadora":("devolucion en centro","recoger devolucion","centro de atencion","devolver importe","devolucion de importe"),
    }
    scores=[]
    for order,(matter,aliases) in enumerate(routing.items()):
        score=0
        for alias in aliases:
            if alias in folded:
                score+=max(2,len(alias.split()))
        matter_tokens=_legal_tokens(matter)
        score+=len(_legal_tokens(obligation)&matter_tokens)
        scores.append((score,-order,matter))
    scores.sort(reverse=True)
    primary=scores[0][2] if scores and scores[0][0]>0 else "Todas las materias"
    selected=[r for r in rows if r["materia"] in {primary,"Todas las materias","Notificaciones"}]
    ids=[int(r["fila"]) for r in selected]
    header=("Fuente consultada: criterios_evaluacion_obligatorios.txt. "
            f"Materia seleccionada desde la obligación principal: {primary}. "
            "Solo las filas siguientes pueden utilizarse para este razonamiento; "
            "las filas de otras materias quedan excluidas para evitar cruces.")
    text=header+"\n"+"\n".join(f"[FILA {r['fila']} / {r['materia']}] {r['texto']}" for r in selected)
    return primary,text,ids

def classify_service_restriction(obligation: str, matter: str) -> tuple[bool,str]:
    """Classify only mandates that directly restore or enable use of the service."""
    text=_fold_legal_text(obligation)
    information_terms=(
        "entregar el contrato","remitir el contrato","informar al reclamante",
        "informar al usuario","terminos y condiciones","montos acordados",
        "tarifas aplicadas","brindar informacion","entregar informacion",
        "remitir informacion","trasladar informacion")
    restoration_terms=(
        "reconect","reactiv","restablec","instalar el servicio",
        "instalacion del servicio","trasladar el servicio",
        "traslado del servicio","retornar el numero","devolver el numero",
        "portabilidad no autorizada","suspension injustificada",
        "baja injustificada")
    if ((any(term in text for term in information_terms) or is_information_delivery_obligation(obligation))
            and not any(term in text for term in restoration_terms)):
        return False, ("La obligación consiste en entregar o comunicar información contractual; "
                       "no restablece ni habilita materialmente el acceso al servicio.")

    restricted=any(term in text for term in restoration_terms)
    restricted=restricted or (
        any(term in text for term in ("averia","interrupcion","falta de servicio","operatividad")) and
        any(term in text for term in ("prueba","verificar","reparar","operativ","restablec")))
    restricted=restricted or (
        "cambio de titularidad" in text and
        any(term in text for term in ("falta de servicio","reactiv","restablec","activar el servicio")))
    if restricted:
        return True, ("La obligación principal restablece o habilita el uso material del servicio "
                      "en una de las materias institucionalmente calificadas como restrictivas.")
    return False, ("La obligación no corresponde a calidad o falta de servicio por avería, retorno "
                   "por portabilidad no autorizada, reactivación por suspensión o baja injustificada, "
                   "instalación, traslado ni falta de servicio por cambio de titularidad.")

def legal_sources(case_context: str, template: str, obligation: str) -> dict[str,str]:
    matter,criteria,criteria_ids=select_applicable_criteria(obligation)
    restriction,restriction_reason=classify_service_restriction(obligation,matter)
    information_delivery=is_information_delivery_obligation(obligation)
    return {
        "instrucciones":INSTRUCCIONES.read_text("utf-8",errors="ignore"),
        "materia_criterios_seleccionada":matter,
        "filas_criterios_seleccionadas":", ".join(map(str,criteria_ids)),
        "criterios_evaluacion_aplicables":criteria,
        "restriccion_servicio":"SI" if restriction else "NO",
        "sustento_restriccion_servicio":restriction_reason,
        "plantilla_aplicable":source_text(template,80000),
        "pautas_pas_aplicables":relevant_excel_rules(
            "PAUTAS PAS.xlsx",case_context,26,service_restriction=restriction,
            information_delivery=information_delivery),
    }

def calculate_due(notification: str, context: str) -> tuple[str,str] | None:
    """Calculate only the principal term already extracted from HA RESUELTO."""
    try:
        normalized=unicodedata.normalize("NFD",context.lower())
        normalized="".join(c for c in normalized if unicodedata.category(c)!="Mn")
        start=parse_excel_date(notification)
        number_match=re.search(r"(?:\(|\b)(\d{1,3})\)?\s*(?:dias?|mes(?:es)?)",normalized,re.I)
        word_numbers={"un":1,"uno":1,"dos":2,"tres":3,"cuatro":4,"cinco":5,"seis":6,
                      "siete":7,"ocho":8,"nueve":9,"diez":10,"quince":15,"veinte":20}
        word_match=re.search(r"\b("+"|".join(word_numbers)+r")\b\s*(?:dias?|mes(?:es)?)",normalized,re.I)
        quantity=int(number_match.group(1)) if number_match else (word_numbers[word_match.group(1)] if word_match else None)
        if quantity is None:
            raise ValueError("el plazo principal no contiene una cantidad identificable")
        if "mes" in normalized:
            label=f"{quantity} mes"+("es" if quantity!=1 else "")
            return (start+pd.DateOffset(months=quantity)).strftime("%d/%m/%Y"),label
        if "calendario" in normalized or "naturales" in normalized:
            return (start.normalize()+pd.Timedelta(days=quantity)).strftime("%d/%m/%Y"),f"{quantity} días calendario"
        if "habil" not in normalized:
            raise ValueError("no se identificó si el plazo principal se computa en días hábiles")
        days=quantity
        term="un (1) día hábil" if days==1 else f"{days} días hábiles"
        holidays_book=pd.read_excel(FUENTES/"CONTADOR DE PLAZOS - TRASU 2026.xlsx",sheet_name="No laborables (2)",header=None)
        if holidays_book.shape[1] < 2: raise ValueError("la hoja No laborables (2) no contiene la columna Lima")
        holidays=pd.to_datetime(holidays_book.iloc[:,1],dayfirst=True,errors="coerce").dropna().dt.normalize().unique()
        offset=pd.offsets.CustomBusinessDay(n=days,weekmask="Mon Tue Wed Thu Fri",holidays=list(holidays))
        return (start.normalize()+offset).strftime("%d/%m/%Y"),term
    except Exception as e:
        raise ValueError(f"error del contador de plazos: {e}") from e

def deterministic_paragraph(result: dict[str,Any], extraction: dict[str,Any]) -> str:
    """Last-resort grounded drafting; never leaves a completed case blank."""
    f=result.get("ficha") if isinstance(result.get("ficha"),dict) else {}
    acto=f.get("tipo_acto") or "acto evaluado"
    notif=f.get("fecha_notificacion_emision") or "no identificada"
    plazo=f.get("plazo_cumplimiento") or "no identificado"
    vence=f.get("fecha_vencimiento") or "no identificada"
    obligacion=f.get("obligacion_principal") or "la obligación ordenada"
    pruebas=[]
    for item in (extraction.get("medios_probatorios") or []):
        if not isinstance(item,dict): continue
        doc=item.get("documento") or "documento aportado"
        fecha=f" de fecha {item.get('fecha')}" if item.get("fecha") else ""
        hecho=item.get("hecho_acreditado") or "sin hecho acreditado identificado"
        estado=item.get("estado") or "no acreditado"
        pruebas.append(f"{doc}{fecha}: {hecho} (estado: {estado})")
    matrices=[]
    for item in (extraction.get("matriz_cumplimiento") or []):
        if isinstance(item,dict):
            matrices.append(f"{item.get('componente') or 'obligación'}: {item.get('estado') or 'no acreditado'}; {item.get('sustento') or 'sin sustento adicional'}")
    contraste="; ".join(pruebas) if pruebas else "no se identificaron medios probatorios objetivos suficientes"
    matriz="; ".join(matrices) if matrices else "no se acreditó documentalmente la ejecución de cada componente"
    resultado=result.get("resultado") or "Pendiente"
    subs=result.get("subsanacion_voluntaria") or "no corresponde"
    clas=result.get("clasificacion") or "Pendiente"
    return (f"El {acto} fue notificado el {notif}, con un plazo de {plazo}, que vencía el {vence}. "
            f"La obligación principal consistía en {obligacion}. Al respecto, la documentación examinada acredita lo siguiente: {contraste}. "
            f"Del contraste individual de las obligaciones se obtiene: {matriz}. En consecuencia, el resultado de la evaluación es {resultado} y la clasificación es {clas}. "
            f"Respecto de la subsanación voluntaria, {subs}, conforme al cese y reversión efectivamente acreditados en el expediente.")

def normalize_legal_paragraph(paragraph: str, ficha: dict[str,Any], resultado: str="") -> str:
    """Enforce template wording and dd/mm/yyyy dates on every generation route."""
    text=str(paragraph or "").strip()
    months={"enero":"01","febrero":"02","marzo":"03","abril":"04","mayo":"05","junio":"06",
            "julio":"07","agosto":"08","septiembre":"09","setiembre":"09","octubre":"10",
            "noviembre":"11","diciembre":"12"}
    pattern=r"\b(\d{1,2})\s+de\s+("+"|".join(months)+r")\s+de\s+(\d{4})\b"
    text=re.sub(pattern,lambda m:f"{int(m.group(1)):02d}/{months[m.group(2).lower()]}/{m.group(3)}",text,flags=re.I)
    text=re.sub(r"\b(\d{4})-(\d{2})-(\d{2})\b",lambda m:f"{m.group(3)}/{m.group(2)}/{m.group(1)}",text)
    is_trasu=("TRASU" in str(ficha.get("tipo_acto") or "").upper() or
              "TRASU" in str(ficha.get("numero_acto") or "").upper() or
              "TRASU" in text[:180].upper())
    if is_trasu:
        notification=str(ficha.get("fecha_notificacion_emision") or "").strip()
        due=str(ficha.get("fecha_vencimiento") or "").strip()
        term=str(ficha.get("plazo_cumplimiento") or "").strip()
        obligation=str(ficha.get("obligacion_principal") or "").strip()
        if obligation:
            obligation=obligation.rstrip(" .;,")
            obligation=obligation[:1].lower()+obligation[1:]
        invalid={"","no identificado","pendiente de verificación","pendiente de verificacion"}
        if notification.lower() not in invalid and due.lower() not in invalid and term.lower() not in invalid:
            # The legal header is deterministic. Gemini may draft only the
            # evidentiary analysis that begins with "Al respecto".
            match=re.search(r"\bAl\s+respecto\b",text,flags=re.I)
            tail=text[match.start():].strip() if match else ""
            opening=(f"La Resolución TRASU fue notificada el {notification}, otorgando a la empresa operadora "
                     f"el plazo de {term} para cumplir la obligación principal consistente en {obligation}, "
                     f"plazo que vencía el {due}.")
            text=opening+(" "+tail if tail else "")
        else:
            raise ValueError("No se puede redactar un párrafo TRASU sin notificación, plazo y vencimiento verificados")
        folded_obligation=unicodedata.normalize("NFD",obligation.lower())
        folded_obligation="".join(c for c in folded_obligation if unicodedata.category(c)!="Mn")
        if any(k in folded_obligation for k in ("reconect","reactiv","restablec")):
            # Final deterministic guard against cross-matter language. Gemini
            # sometimes copies discount-specific wording even after the matter
            # filter. Such clauses are impossible in a reconnection analysis.
            text=re.sub(
                r"[,;]\s*(?:al\s+haber\s+)?(?:quedando|existiendo)?\s*(?:periodos?|meses?)\b[^.;]*?(?:descuent\w*|benefici\w*|importes?|montos?|registro\s+y\s+activaci[oó]n)[^.;]*[;.]?",
                ";",text,flags=re.I)
            text=re.sub(
                r",\s*toda\s+vez\s+que,?\s*(?:al\s+)?(?:haber\s+)?(?:quedar|quedando|existir|existiendo)\s+(?:periodos?|meses?)\b[^.]*?registro\s+y\s+activaci[oó]n[^.]*\.",
                ", toda vez que no existe una consulta o registro histórico con fecha objetiva que acredite que el servicio ya estaba activo o fue reconectado dentro del plazo otorgado.",
                text,flags=re.I)
            text=re.sub(
                r"[,;]\s*(?:al\s+)?(?:haber\s+)?(?:quedar|quedando|existir|existiendo)\s+(?:periodos?|meses?)\b[^.]*?registro\s+y\s+activaci[oó]n[^.]*[.;]?",
                ";",text,flags=re.I)
            text=re.sub(r";\s*;",";",text)
            text=re.sub(r"\s+([,;:.])",r"\1",text)
            # Sentence-level fallback: reject the whole contaminated sentence,
            # regardless of whether Gemini writes "quedando", "al quedar" or
            # another grammatical variant.
            clean_sentences=[]
            for sentence in re.split(r"(?<=[.!?])\s+",text):
                folded_sentence=unicodedata.normalize("NFD",sentence.lower())
                folded_sentence="".join(c for c in folded_sentence if unicodedata.category(c)!="Mn")
                wrong_period_language=(
                    any(k in folded_sentence for k in ("periodo","meses")) and
                    any(k in folded_sentence for k in ("registro y activacion","descuento","beneficio recurrente","nota de credito","importe","monto")))
                if not wrong_period_language: clean_sentences.append(sentence)
            text=" ".join(clean_sentences)
    # A completed paragraph must never jump from an empty "En consecuencia"
    # directly to subsanation. Restore the missing conclusion from the locked
    # evaluation result without asking the model to redraft anything.
    result_key=unicodedata.normalize("NFD",str(resultado or "").lower())
    result_key="".join(c for c in result_key if unicodedata.category(c)!="Mn")
    if result_key=="incumplio":
        conclusion="En consecuencia, al no haberse acreditado la ejecución oportuna del mandato, se habría configurado la infracción. Asimismo"
    elif result_key=="cumplio":
        conclusion="En consecuencia, se acreditó el cumplimiento oportuno del mandato. Asimismo"
    elif result_key=="inejecutable":
        conclusion="En consecuencia, se determinó que el mandato era inejecutable. Asimismo"
    else:
        conclusion="En consecuencia, no fue posible establecer una conclusión definitiva. Asimismo"
    text=re.sub(r"\bEn\s+consecuencia\s*[;,.]?\s*Asimismo\b",conclusion,text,flags=re.I)
    # Some otherwise valid drafts omit the conclusion entirely and jump from
    # the evidentiary contrast directly to subsanation. Insert only the locked
    # conclusion; preserve every preceding fact and every following analysis.
    if not re.search(r"\bEn\s+consecuencia\b",text,flags=re.I):
        text=re.sub(r"\bAsimismo\b",conclusion,text,count=1,flags=re.I)
    return text

def enforce_conditional_reconnection_timeline(result: dict[str,Any], extraction: dict[str,Any],
                                              paragraph: str, sources: dict[str,str],
                                              documents: dict[str,str] | None=None) -> str:
    """Reject the inference that a later active state proves the earlier condition or timely execution."""
    ficha=result.get("ficha") if isinstance(result.get("ficha"),dict) else {}
    obligation=_fold_legal_text(ficha.get("obligacion_principal"))
    selected_matter=str(sources.get("materia_criterios_seleccionada") or "")
    conditional_reconnection=(
        selected_matter=="Corte o Baja Injustificada" and
        any(k in obligation for k in ("reconect","reactiv","restablec")) and
        any(k in obligation for k in ("si ","suspend")))
    if not conditional_reconnection:
        return paragraph
    try:
        notification=parse_excel_date(ficha.get("fecha_notificacion_emision")).normalize()
        due=parse_excel_date(ficha.get("fecha_vencimiento")).normalize()
    except Exception:
        return paragraph

    timely_objective_proof=False
    timely_proof_state=""
    timely_proof_date=None
    late_state_dates=[]
    has_system_printer=False
    letter_dates=[]
    for item in (extraction.get("medios_probatorios") or []):
        if not isinstance(item,dict): continue
        state=_fold_legal_text(item.get("estado"))
        date_text=str(item.get("fecha_ejecucion_acreditada") or "").strip()
        source_text=str(item.get("fuente_fecha_ejecucion") or item.get("cita") or "").strip()
        evidence_text=_fold_legal_text(" ".join(str(item.get(k) or "") for k in
                                                ("documento","naturaleza","hecho_acreditado","cita","fuente_fecha_ejecucion")))
        document_text=_fold_legal_text(" ".join(str(item.get(k) or "") for k in
                                                ("documento","naturaleza","cita")))
        # The date of the letter that remitted an undated system printer is
        # evidence metadata. Read it from the extraction, never from the exact
        # wording Gemini happens to use in the final paragraph.
        if ("alegacion" in document_text or "carta" in document_text or
                "rf-t-fc" in document_text or "rf-c-fc" in document_text):
            for field in ("fecha_documento","cita","documento"):
                for extracted_date in _legal_dates_in(item.get(field)):
                    if extracted_date>=notification and extracted_date not in letter_dates:
                        letter_dates.append(extracted_date)
        if any(k in evidence_text for k in ("printer","captura del sistema","consulta del sistema",
                                             "historico del sistema","estado del servicio")):
            has_system_printer=True
        shows_active=("servicio" in evidence_text and any(k in evidence_text for k in ("activo","operativo")))
        if state in {"ejecutado","condicion_no_configurada"} and date_text and source_text:
            event_date=pd.to_datetime(date_text,dayfirst=True,errors="coerce")
            if pd.notna(event_date):
                event_date=pd.Timestamp(event_date).normalize()
                if state=="condicion_no_configurada" and event_date==notification:
                    timely_objective_proof=True; timely_proof_state=state; timely_proof_date=event_date; break
                if state=="ejecutado" and notification<=event_date<=due:
                    timely_objective_proof=True; timely_proof_state=state; timely_proof_date=event_date; break
        if shows_active:
            # A system printer without its own visible date is not dated with
            # the enclosing letter's date. The letter date is handled below as
            # the reference date for remitting an otherwise undated proof.
            later_text=str(item.get("fecha_ejecucion_acreditada") or "").strip()
            later_date=pd.to_datetime(later_text,dayfirst=True,errors="coerce")
            if source_text and pd.notna(later_date) and pd.Timestamp(later_date).normalize()>due:
                late_state_dates.append(pd.Timestamp(later_date).normalize())
    # Gemini can omit an attached printer from its JSON even though it is
    # expressly inventoried in the source documents. The deterministic guard
    # therefore checks the actual extracted case text as well. This does not
    # invent a date: it only confirms that system evidence was in fact supplied.
    raw_documents="\n".join(str(value or "") for value in (documents or {}).values())
    raw_folded=_fold_legal_text(raw_documents)
    raw_has_printer=any(term in raw_folded for term in (
        "printer","captura del sistema","consulta del sistema","historico del sistema",
        "historico de estado","estado de cuenta financiera","recibos ajustados"))
    raw_reports_active=("servicio" in raw_folded and
                        any(term in raw_folded for term in ("activo","operativo","sin suspension vigente",
                                                            "no registrando suspension vigente")))
    if raw_has_printer and raw_reports_active:
        has_system_printer=True
    if timely_objective_proof:
        if timely_proof_state=="condicion_no_configurada":
            ficha["estado_ejecucion"]="No aplica"
            ficha["fecha_ejecucion"]="No aplica"
        else:
            ficha["estado_ejecucion"]="Ejecutó"
            ficha["fecha_ejecucion"]=timely_proof_date.strftime("%d/%m/%Y") if timely_proof_date is not None else "No identificada"
        return paragraph

    result["resultado"]="Incumplió"
    result["clasificacion"]="PAS"
    result["subsanacion_voluntaria"]="no aplica"
    checklist=result.setdefault("evaluacion_juridica",{}).setdefault("checklist",[])
    respect=re.search(r"\bAl\s+respecto\b.*?\.(?=\s+[A-ZÁÉÍÓÚÑ]|$)",str(paragraph or ""),flags=re.I|re.S)
    allegation=respect.group(0).strip() if respect else (
        "Al respecto, la empresa operadora señaló que el servicio se encontraba activo y operativo.")
    # A letter contains the operator's allegation; it must not be described as
    # the objective proof when the actual verification comes from its attached
    # system printers or screenshots.
    allegation=re.sub(
        r",?\s*adjuntando\s+(?:el\s+)?documento\s+RF[-_A-Z0-9.]+(?:\.pdf)?\s+como\s+sustento",
        "",allegation,flags=re.I)
    allegation=re.sub(r"\s+([,.])",r"\1",allegation)
    letter_date=None
    letter_match=re.search(r"\bcarta\s+de\s+fecha\s+(\d{1,2}/\d{1,2}/\d{4})",allegation,flags=re.I)
    if letter_match:
        parsed_letter=pd.to_datetime(letter_match.group(1),dayfirst=True,errors="coerce")
        if pd.notna(parsed_letter): letter_date=pd.Timestamp(parsed_letter).normalize()
    if letter_date is None and letter_dates:
        # Use the earliest post-deadline remitting letter as the conservative
        # reference. If none is late, use the earliest available letter.
        post_deadline=sorted(d for d in letter_dates if d>due)
        letter_date=(post_deadline or sorted(letter_dates))[0]
    analysis_start=re.search(r"\bAl\s+respecto\b",str(paragraph or ""),flags=re.I)
    preserved_opening=str(paragraph or "")[:analysis_start.start()].strip() if analysis_start else ""
    def with_preserved_opening(analysis: str) -> str:
        return (preserved_opening+" "+analysis).strip() if preserved_opening else analysis
    notification_text=notification.strftime("%d/%m/%Y")
    due_text=due.strftime("%d/%m/%Y")
    undated_printer_with_letter=(has_system_printer and not late_state_dates and
                                letter_date is not None and letter_date>due)
    if late_state_dates or undated_printer_with_letter:
        accredited_timestamp=min(late_state_dates) if late_state_dates else letter_date
        accredited_date=accredited_timestamp.strftime("%d/%m/%Y")
        ficha["estado_ejecucion"]="Ejecutó"
        ficha["fecha_ejecucion"]=accredited_date
        proof_subject=("los printers de los sistemas de la empresa operadora adjuntos a la carta"
                       if has_system_printer else "los medios probatorios objetivos adjuntos a la carta")
        checklist.append(
            f"Fuente probatoria: {proof_subject}; el estado activo se acredita recién el {accredited_date}, después del "
            f"vencimiento del {due_text}; corresponde ejecución fuera de plazo.")
        if undated_printer_with_letter:
            evidentiary_sentence=(
                f"Ahora bien, de la revisión de {proof_subject}, se verifica que el servicio se encontraba activo y "
                f"operativo; dado que dichos printers no consignan una fecha propia, se toma como fecha de referencia "
                f"la carta que los remitió, esto es, el {accredited_date}. No obstante, esa fecha es posterior al "
                f"vencimiento del {due_text}.")
        else:
            evidentiary_sentence=(
                f"Ahora bien, de la revisión de {proof_subject}, se verifica que el servicio se encontraba activo y "
                f"operativo el {accredited_date}; sin embargo, no contiene un registro histórico con una fecha anterior "
                f"que demuestre la ejecución hasta el {due_text}.")
        return with_preserved_opening(
                f"{allegation} {evidentiary_sentence} En consecuencia, la empresa operadora acreditó la ejecución del "
                "mandato fuera del plazo establecido, con lo cual se habría configurado la infracción. Asimismo, no "
                "resulta aplicable el eximente de subsanación voluntaria, debido a que el mandato materia de análisis "
                "se encuentra vinculado a una restricción del servicio cuyos efectos no pueden ser revertidos.")
    ficha["estado_ejecucion"]="No ejecutó"
    ficha["fecha_ejecucion"]="No aplica"
    checklist.append(
        "Validación temporal: no existe prueba objetiva fechada que acredite que la condición no se configuró "
        "en la fecha de notificación o que la reconexión se ejecutó hasta el vencimiento.")
    return with_preserved_opening(
            f"{allegation} Ahora bien, dicho medio acredita únicamente el estado del servicio en la fecha de "
            f"su consulta o comunicación, pero no contiene un registro histórico con fecha objetiva que demuestre que "
            f"el servicio ya estaba activo el {notification_text} ni que hubiese sido reconectado hasta el {due_text}. "
            "Por tanto, un estado posterior no permite inferir que la condición del mandato no se configuró al momento "
            "de la notificación. En consecuencia, al no haberse acreditado la ejecución oportuna del mandato, se habría "
            "configurado la infracción. Asimismo, no resulta aplicable el eximente de subsanación voluntaria, debido a "
            "que no se acreditó el cese total de la conducta ni la reversión integral de sus efectos.")

def evaluate_subsanacion_especifica(result: dict[str,Any], extraction: dict[str,Any],
                                    sources: dict[str,str] | None) -> tuple[str,str]:
    """Decide subsanación with the selected materia rules only; ('','') when unavailable."""
    try:
        subs_context={"ficha":result.get("ficha",{}),"extraccion_probatoria":extraction,
                      "resultado":result.get("resultado"),
                      "criterios":(sources or {}).get("criterios_evaluacion_aplicables",""),
                      "pautas":(sources or {}).get("pautas_pas_aplicables",""),
                      "restriccion_servicio":(sources or {}).get("restriccion_servicio",""),
                      "sustento_restriccion_servicio":(sources or {}).get("sustento_restriccion_servicio","")}
        raw=parse_json_response(gemini_text(
            "Evalúa únicamente si aplica el eximente de subsanación voluntaria, siguiendo los supuestos de "
            "los criterios y pautas entregados para la materia. El eximente exige conjuntamente cese total, "
            "reversión integral y oportunidad anterior al inicio del PAS. Si restriccion_servicio=SI, el "
            "eximente no procede por la irreversibilidad de la falta de servicio; si es NO, está prohibido "
            "negarlo por efectos irreversibles de una restricción y deben evaluarse separadamente el cese "
            "total, la reversión integral y su oportunidad. Devuelve JSON "
            "{subsanacion_voluntaria:'aplica'|'no aplica','sustento':string}.",
            json.dumps(subs_context,ensure_ascii=False),json_mode=True))
        candidate=_fold_legal_text((raw or {}).get("subsanacion_voluntaria") if isinstance(raw,dict) else "")
        sustento=str(raw.get("sustento","")).strip() if isinstance(raw,dict) else ""
        if "no aplica" in candidate: return "no aplica",sustento
        if "aplica" in candidate: return "aplica",sustento
    except Exception:
        pass
    return "",""

def enforce_service_restriction_rule(result: dict[str,Any], paragraph: str,
                                     sources: dict[str,str],
                                     extraction: dict[str,Any] | None=None) -> str:
    """Prevent a lack-of-service PAS rule from contaminating other matters."""
    restricted=str(sources.get("restriccion_servicio") or "NO").upper()=="SI"
    reason=str(sources.get("sustento_restriccion_servicio") or "").strip()
    # Regla transversal (fila 83): entregar contratos o información nunca es
    # una restricción del servicio, sin importar cómo se haya clasificado antes.
    obligation=str((result.get("ficha") or {}).get("obligacion_principal") or "")
    if restricted and is_information_delivery_obligation(obligation):
        restricted=False
        reason=("Corrección transversal: la obligación consiste en entregar o comunicar información "
                "contractual; no restablece ni habilita materialmente el acceso al servicio.")
    evaluation=result.setdefault("evaluacion_juridica",{})
    evaluation["restriccion_servicio"]="Sí" if restricted else "No"
    checklist=evaluation.setdefault("checklist",[])
    note=(f"Clasificación de restricción del servicio: {'Sí' if restricted else 'No'}. {reason}").strip()
    if note not in checklist:
        checklist.append(note)
    if restricted:
        return paragraph

    sentences=re.split(r"(?<=[.!?])\s+",str(paragraph or "").strip())
    def _is_false_restriction(sentence: str) -> bool:
        folded=_fold_legal_text(sentence)
        return ("restriccion del servicio" in folded and
                any(term in folded for term in ("no pueden ser revertidos","no puede ser revertido",
                                                "efectos no pueden","no resulta aplicable")))
    # Si la única razón para negar la subsanación fue la restricción falsa, la
    # conclusión quedó sin fundamento: reevaluarla con las reglas de la materia.
    if (any(_is_false_restriction(s) for s in sentences) and
            _fold_legal_text(result.get("subsanacion_voluntaria"))=="no aplica" and
            "incumpl" in _fold_legal_text(result.get("resultado"))):
        value,sustento=evaluate_subsanacion_especifica(result,extraction or {},sources)
        if value:
            result["subsanacion_voluntaria"]=value
            result["clasificacion"]="NO PAS" if value=="aplica" else "PAS"
            if sustento:
                result.setdefault("evaluacion_juridica",{}).setdefault("checklist",[]).append(
                    "Subsanación voluntaria reevaluada al descartarse la restricción del servicio: "+sustento)

    subs=_fold_legal_text(result.get("subsanacion_voluntaria"))
    if subs=="aplica":
        replacement=("Asimismo, resulta aplicable el eximente de subsanación voluntaria, conforme al cese total, "
                     "la reversión integral y la oportunidad anterior al inicio del PAS acreditados en la evaluación.")
    elif subs=="no corresponde":
        replacement="Asimismo, no corresponde evaluar subsanación voluntaria en este supuesto."
    else:
        replacement=("Asimismo, no resulta aplicable el eximente de subsanación voluntaria, debido a que no se "
                     "acreditaron conjuntamente el cese total, la reversión integral y su oportunidad anterior al "
                     "inicio del PAS.")
    cleaned=[]
    replaced=False
    for sentence in sentences:
        if _is_false_restriction(sentence):
            if not replaced:
                cleaned.append(replacement); replaced=True
            continue
        cleaned.append(sentence)
    if not replaced and "incumpl" in _fold_legal_text(result.get("resultado")):
        cleaned.append(replacement)
    return " ".join(x for x in cleaned if x).strip()

def enforce_information_delivery_language(result: dict[str,Any], paragraph: str,
                                          documents: dict[str,str] | None=None) -> str:
    """Keep contract/information cases free of discount rules and invented server events."""
    ficha=result.get("ficha") if isinstance(result.get("ficha"),dict) else {}
    obligation=str(ficha.get("obligacion_principal") or "")
    if not is_information_delivery_obligation(obligation):
        return paragraph
    raw=_fold_legal_text("\n".join(f"{name}\n{text}" for name,text in (documents or {}).items()))
    paragraph_folded=_fold_legal_text(paragraph)
    email_case="correo" in (raw+" "+paragraph_folded)
    arrival_confirmed=any(x in raw for x in (
        "constancia de entrega","correo entregado","mensaje entregado",
        "entrega exitosa","entregado al destinatario","delivered",
        "constancia de recepcion","acuse de recibo","recepcion del usuario")) or any(
        x in paragraph_folded for x in (
            "la entrega del contrato se realizo","el contrato fue entregado",
            "el correo fue entregado","se acredito la entrega"))
    checklist=result.setdefault("evaluacion_juridica",{}).setdefault("checklist",[])
    note=("Control de entrega de información: se aplicó exclusivamente la prueba de contenido, envío, "
          "entrega y recepción; se excluyeron reglas de periodos, descuentos, registro y activación.")
    if note not in checklist:
        checklist.append(note)
    cleaned=[]
    neutral_added=False
    obligation_folded=_fold_legal_text(obligation)
    mandate_has_period=any(x in obligation_folded for x in ("periodo","meses","mensual"))
    for sentence in re.split(r"(?<=[.!?])\s+",str(paragraph or "").strip()):
        folded=_fold_legal_text(sentence)
        cross_matter=(
            (any(x in folded for x in ("periodo","meses pendientes","meses restantes")) and
             (not mandate_has_period or
              any(x in folded for x in ("registro y activacion","descuento","beneficio","nota de credito")))) or
            (any(x in folded for x in ("registro y activacion","descuento recurrente","nota de credito")) and
             not any(x in obligation_folded for x in ("registro","descuento","nota de credito"))))
        if cross_matter:
            if not neutral_added:
                cleaned.append(
                    "Al no haberse acreditado la entrega y recepción de todos los documentos y de toda la "
                    "información ordenada, no se acreditó la ejecución íntegra del mandato.")
                neutral_added=True
            continue
        invented_server=(
            "servidor" in folded and
            any(x in folded for x in ("no envio","no genero","no confirmo","no remitio")) and
            not ("servidor" in raw and any(x in raw for x in ("no envio","no genero","no confirmo","no remitio"))))
        if invented_server:
            cleaned.append(
                "La ausencia de una constancia no permite atribuir al servidor una respuesta que no figura "
                "literalmente en el expediente; únicamente permite señalar que no obra una confirmación "
                "documental de entrega y recepción con fecha verificable.")
            continue
        improper_additional_receipt=(
            email_case and
            any(x in folded for x in (
                "recepcion efectiva por parte del usuario","confirmacion de lectura",
                "respuesta del usuario","acuse de recibo",
                "confirmacion de recepcion")) and
            any(x in folded for x in ("no existe","no obra","no acredita","sin constancia")) and
            arrival_confirmed)
        if improper_additional_receipt:
            due_dates=_legal_dates_in(ficha.get("fecha_vencimiento"))
            sentence_dates=_legal_dates_in(sentence)
            if arrival_confirmed and sentence_dates:
                execution_date=max(sentence_dates)
                due_date=due_dates[0] if due_dates else None
                timing=(
                    f"; sin embargo, dicha ejecución fue posterior al vencimiento del {due_date.strftime('%d/%m/%Y')}"
                    if due_date is not None and execution_date>due_date else
                    (f" y se produjo dentro del plazo que vencía el {due_date.strftime('%d/%m/%Y')}"
                     if due_date is not None else ""))
                cleaned.append(
                    "La constancia de envío, acompañada por la constancia de entrega al destinatario o, "
                    "alternativamente, por la constancia de recepción del usuario, acredita que el correo "
                    f"electrónico fue entregado el {execution_date.strftime('%d/%m/%Y')}{timing}.")
            elif arrival_confirmed:
                cleaned.append(
                    "La constancia de envío, acompañada por la constancia de entrega al destinatario o, "
                    "alternativamente, por la constancia de recepción del usuario, acredita la entrega del correo "
                    "electrónico; no se exige acumular ambos medios.")
            continue
        cleaned.append(sentence)
    text=" ".join(x for x in cleaned if x).strip()
    # Catch the frequent contamination even when it appears inside a longer sentence.
    text=re.sub(
        r"(?:Al\s+)?(?:quedar|quedando|existir|existiendo)\s+(?:periodos?|meses?)\b[^.]*?"
        r"(?:registro\s+y\s+activaci[oó]n|descuent\w*|benefici\w*|nota\s+de\s+cr[eé]dito)[^.]*\.?",
        "",text,flags=re.I)
    text=re.sub(r"\s+([,;:.])",r"\1",text)
    text=re.sub(r"\s{2,}"," ",text).strip()
    return text

_FILE_EXTENSION_RE=re.compile(r"\.(pdf|tiff?|jpe?g|png|bmp|docx?|xlsx?|csv|zip|rar|7z)\b",re.I)

def enforce_evidence_citation(extraction: dict[str,Any], paragraph: str) -> str:
    """The model does not reliably follow the instruction to cite each evidence
    document by name and date (temperature is not zero, so compliance varies
    run to run). Guarantee it deterministically instead of relying on the
    prompt alone: append any objective evidence document the model omitted.
    Documents are identified by the denomination visible in their content;
    computer file names never identify a document, so a documento value that
    looks like a file name is skipped and stray extensions are stripped."""
    items=[x for x in (extraction.get("medios_probatorios") or []) if isinstance(x,dict)]
    paragraph=_FILE_EXTENSION_RE.sub("",str(paragraph or "").strip())
    if not paragraph or not items:
        return paragraph
    folded_paragraph=_fold_legal_text(paragraph)
    missing=[]; seen=set()
    for item in items:
        doc=str(item.get("documento") or "").strip()
        date=str(item.get("fecha_ejecucion_acreditada") or item.get("fecha_documento") or "").strip()
        if not doc or not date or _FILE_EXTENSION_RE.search(doc):
            continue
        key=(doc.lower(),date)
        if key in seen:
            continue
        seen.add(key)
        if _fold_legal_text(doc) in folded_paragraph and date in paragraph:
            continue
        missing.append(f"{doc} de fecha {date}")
    if not missing:
        return paragraph
    addition="Al respecto, obran en el expediente los siguientes medios probatorios: "+"; ".join(missing)+"."
    match=re.search(r"\bEn\s+consecuencia\b",paragraph)
    if match:
        return (paragraph[:match.start()].rstrip()+" "+addition+" "+paragraph[match.start():]).strip()
    return (paragraph+" "+addition).strip()

def normalize_verdict_fields(result: dict[str,Any], extraction: dict[str,Any],
                             sources: dict[str,str] | None=None) -> None:
    """The dictamen uses a closed vocabulary: resultado ∈ {Cumplió, Incumplió,
    Inejecutable}, subsanación ∈ {aplica, no aplica, no corresponde} and
    clasificación ∈ {PAS, NO PAS}. 'Pendiente' must never reach the UI, so any
    empty or off-vocabulary value is normalized here. A missing subsanación is
    never assumed from the resultado alone: it is taken from the drafted
    paragraph or, failing that, evaluated against the selected criterios and
    pautas, because each materia has its own subsanación supuestos."""
    folded=_fold_legal_text(result.get("resultado"))
    if "incumpl" in folded: resultado="Incumplió"
    elif "inejecut" in folded: resultado="Inejecutable"
    elif "cumpl" in folded: resultado="Cumplió"
    else:
        evidence={_fold_legal_text(x.get("estado")) for x in (extraction.get("medios_probatorios") or []) if isinstance(x,dict)}
        matrix={_fold_legal_text(x.get("estado")) for x in (extraction.get("matriz_cumplimiento") or []) if isinstance(x,dict)}
        paragraph=_fold_legal_text(result.get("parrafo_final"))
        if (evidence & {"programado","en_curso","no_acreditado"} or matrix & {"parcial","no_acreditado"}
                or "configurado la infraccion" in paragraph or "incumpl" in paragraph):
            resultado="Incumplió"
        elif "inejecutable" in paragraph:
            resultado="Inejecutable"
        elif matrix and matrix<={"acreditado"}:
            resultado="Cumplió"
        else:
            # Sin matriz ni estados: la evaluación no acreditó la ejecución íntegra.
            resultado="Incumplió"
    result["resultado"]=resultado
    subs=_fold_legal_text(result.get("subsanacion_voluntaria"))
    if "no aplica" in subs: subs_value="no aplica"
    elif "no corresponde" in subs: subs_value="no corresponde"
    elif "aplica" in subs: subs_value="aplica"
    else:
        subs_value=""
        # 1) El párrafo redactado ya contiene la conclusión de subsanación.
        paragraph=_fold_legal_text(result.get("parrafo_final"))
        if "no corresponde evaluar subsanacion" in paragraph: subs_value="no corresponde"
        elif re.search(r"no\s+(resulta|es)\s+aplicable\s+el\s+eximente",paragraph): subs_value="no aplica"
        elif re.search(r"(resulta|es)\s+aplicable\s+el\s+eximente",paragraph): subs_value="aplica"
        if not subs_value and resultado!="Incumplió":
            # Solo se evalúa subsanación cuando hubo incumplimiento.
            subs_value="no corresponde"
        if not subs_value:
            # 2) Componentes pendientes: por regla transversal no existe cese
            # total ni reversión integral, requisito conjunto del eximente.
            evidence={_fold_legal_text(x.get("estado")) for x in (extraction.get("medios_probatorios") or []) if isinstance(x,dict)}
            matrix={_fold_legal_text(x.get("estado")) for x in (extraction.get("matriz_cumplimiento") or []) if isinstance(x,dict)}
            if evidence & {"programado","en_curso","no_acreditado"} or matrix & {"parcial","no_acreditado"}:
                subs_value="no aplica"
        if not subs_value:
            # 3) Evaluación específica con los criterios y pautas de la materia:
            # cada materia tiene sus propios supuestos de subsanación y no puede
            # asumirse un valor a partir del solo resultado.
            subs_value,sustento=evaluate_subsanacion_especifica(result,extraction,sources)
            if subs_value and sustento:
                result.setdefault("evaluacion_juridica",{}).setdefault("checklist",[]).append(
                    f"Subsanación voluntaria (evaluación específica): {sustento}")
        if not subs_value:
            # Último recurso trazable: sin prueba de cese total y reversión
            # integral acreditados, el eximente no puede tenerse por configurado.
            subs_value="no aplica"
            result.setdefault("evaluacion_juridica",{}).setdefault("checklist",[]).append(
                "Subsanación voluntaria: el modelo no la determinó y no obra en la evaluación prueba de cese total "
                "y reversión integral anteriores al PAS; se registra 'no aplica' por falta de acreditación.")
    result["subsanacion_voluntaria"]=subs_value
    clas=_fold_legal_text(result.get("clasificacion"))
    if "no pas" in clas: clas_value="NO PAS"
    elif "pas" in clas: clas_value="PAS"
    else: clas_value="PAS" if resultado=="Incumplió" and subs_value!="aplica" else "NO PAS"
    result["clasificacion"]=clas_value

def derive_execution_fields(result: dict[str,Any], extraction: dict[str,Any]) -> None:
    """Summarize proven execution without changing the legal result or timeliness analysis."""
    ficha=result.setdefault("ficha",{})
    if str(ficha.get("estado_ejecucion") or "").strip():
        return
    evidence=[x for x in (extraction.get("medios_probatorios") or []) if isinstance(x,dict)]
    matrix=[x for x in (extraction.get("matriz_cumplimiento") or []) if isinstance(x,dict)]
    evidence_states={_fold_legal_text(x.get("estado")) for x in evidence}
    matrix_states={_fold_legal_text(x.get("estado")) for x in matrix}
    dates=[]
    for item in evidence:
        raw=str(item.get("fecha_ejecucion_acreditada") or "").strip()
        parsed=pd.to_datetime(raw,dayfirst=True,errors="coerce")
        if raw and pd.notna(parsed):
            value=pd.Timestamp(parsed).strftime("%d/%m/%Y")
            if value not in dates: dates.append(value)
    has_executed=("ejecutado" in evidence_states or "acreditado" in matrix_states)
    has_pending=bool(evidence_states & {"programado","en_curso","no_acreditado"} or
                     matrix_states & {"parcial","no_acreditado"})
    if "condicion_no_configurada" in evidence_states and not has_executed:
        status="No aplica"
    # La obligación se informa como un todo: cualquier componente pendiente
    # significa que la obligación principal no fue ejecutada íntegramente.
    elif has_pending:
        status="No ejecutó"
    elif has_executed:
        status="Ejecutó"
    elif str(result.get("resultado") or "").strip()=="Incumplió":
        status="No ejecutó"
    else:
        status="Pendiente de verificación"
    ficha["estado_ejecucion"]=status
    if status=="Ejecutó":
        ficha["fecha_ejecucion"]="; ".join(dates) if dates else "No identificada"
    elif status in {"No aplica","No ejecutó"}:
        ficha["fecha_ejecucion"]="No aplica"
    else:
        ficha["fecha_ejecucion"]="Pendiente de verificación"

def ai_evaluate(payload: dict[str,Any]) -> dict[str,Any]:
    if payload.get("tipo_acto")=="Resolución TRASU":
        verified_date=str(payload.get("fecha_verificada") or "").strip()
        verified_due=str(payload.get("fecha_vencimiento") or "").strip()
        invalid={"","no identificado","pendiente de verificación","pendiente de verificacion"}
        if verified_date.lower() in invalid:
            raise ValueError("No se puede evaluar cumplimiento TRASU sin fecha de notificación verificada")
        if verified_due.lower() in invalid:
            raise ValueError("No se puede evaluar cumplimiento TRASU sin fecha de vencimiento calculada")
    template="PLANTILLAS cumplimiento.docx" if payload["tipo_acto"]=="Resolución TRASU" else "PLANTILLAS DENUNCIAS ACTUALIZADAS.docx"
    case_context=json.dumps(payload,ensure_ascii=False)
    verified_principal=payload.get("obligacion_y_plazo_principales_verificados") or {}
    obligation_for_criteria=(str(verified_principal.get("obligacion_principal") or "").strip()
                             if isinstance(verified_principal,dict) else "")
    if not obligation_for_criteria:
        obligation_for_criteria=str(payload.get("parte_resolutiva_trasu") or "").strip()
    sources=legal_sources(case_context,template,obligation_for_criteria)
    extraction_schema={"acto":{"numero":"","fecha":"","mandato_textual":""},"obligacion_extraida_parte_resolutiva":{"texto":"","articulos_numerales":[],"plazo_principal_textual":""},"obligaciones_accesorias_excluidas":[{"texto":"","plazo":"","motivo_exclusion":""}],"obligaciones":[{"componente":"","periodo":"","plazo_expreso":"","prueba_exigible":""}],"medios_probatorios":[{"documento":"","naturaleza":"alegacion|medio_objetivo","fecha_documento":"","fecha_ejecucion_acreditada":"","fuente_fecha_ejecucion":"","hecho_acreditado":"","cita":"","estado":"ejecutado|condicion_no_configurada|programado|en_curso|no_acreditado"}],"matriz_cumplimiento":[{"componente":"","estado":"acreditado|parcial|no_acreditado","sustento":""}],"datos_no_identificados":[]}
    extraction_system="""Actúa como extractor jurídico OSIPTEL. Separa hechos de conclusiones.

REGLA DE ORIGEN DE LA OBLIGACIÓN: si el acto es una Resolución TRASU, lee COMPLETA y CONJUNTAMENTE todos los artículos y numerales del campo parte_resolutiva_trasu. Identifica la obligación material principal relacionada con la prestación o controversia del servicio y su propio plazo, cualquiera sea el verbo o la forma jurídica empleados. La obligación y su plazo pueden estar en numerales consecutivos y no tienen que aparecer en el mismo numeral que declara fundado el reclamo. No uses listas cerradas de verbos. Copia el mandato con fidelidad y sepáralo por componentes, periodos, montos y condiciones.

REGLA OBLIGATORIA — VARIOS NUMERALES EN LA PARTE RESOLUTIVA: vincula cada plazo con la obligación concreta a la que corresponde. Evalúa solamente la ejecución material que resuelve la controversia. Excluye de la evaluación los numerales posteriores que ordenan comunicar o informar al usuario sobre el cumplimiento, informar o acreditar ante el TRASU, remitir constancias o realizar otro reporte formal, junto con sus respectivos plazos. Regístralos en obligaciones_accesorias_excluidas para conservar trazabilidad, pero nunca los mezcles con la obligación o el plazo principales ni los uses como fundamento de incumplimiento.

Los antecedentes y considerandos solo pueden aclarar una referencia o el alcance de un mandato ya contenido en HA RESUELTO. Nunca pueden crear, sustituir o ampliar la obligación. Las alegaciones, cartas posteriores y pruebas de ejecución tampoco pueden definirla.

Para cada prueba distingue ejecución efectiva de solicitud, programación, caso abierto o gestión en curso. Una afirmación de la empresa no prueba por sí sola el hecho. Cita el documento que respalda cada dato.

DENOMINACIÓN DE LOS MEDIOS PROBATORIOS: en el campo documento escribe la denominación que el propio documento muestra en su contenido visible (tipo y número si lo tiene, p. ej. "nota de crédito SQAX-0000097839", "carta RMA-FC384918-2026-AC-1", "printer del sistema", "constancia de entrega"). Está PROHIBIDO usar el nombre del archivo informático o su extensión (.pdf, .tif, .jpg, etc.): el nombre del archivo es irrelevante y puede ser cualquiera; la identidad del documento sale únicamente de su contenido. Si el contenido no muestra denominación ni número, usa solo el tipo genérico visible ("nota de crédito", "captura de pantalla") sin inventar identificadores.

PROHIBICIÓN DE INVENTAR CAUSAS O CALIFICACIONES: en hecho_acreditado y cita, usa exclusivamente los hechos y palabras que el documento fuente expresa literalmente. Está prohibido calificar, categorizar o etiquetar la razón de un pendiente con expresiones técnicas o jurídicas que el documento no contenga literalmente (por ejemplo "restricción sistémica", "limitación técnica", "caso fortuito", "fuerza mayor"), aunque parezcan un resumen razonable. Si el documento dice que un ajuste está "en gestión" o "programado en el sistema", usa exactamente esas palabras o una paráfrasis mínima y fiel; nunca acuñes una etiqueta nueva que no conste en el texto original.

SEPARACIÓN ENTRE ALEGACIÓN Y PRUEBA OBJETIVA: una carta de descargos, como los documentos RF-T-FC, acredita la fecha y el contenido de lo afirmado por la empresa, pero no acredita por sí sola la ejecución material. Registra la carta como naturaleza=alegacion. Registra separadamente como naturaleza=medio_objetivo los printers, capturas, consultas, históricos, recibos, notas de crédito, actas u otros anexos provenientes de los sistemas o actuaciones de la empresa. El razonamiento debe decir que la empresa señaló el hecho mediante la carta y que este se verifica —cuando corresponda— en el medio objetivo adjunto; nunca denomines a la propia carta como el sustento objetivo de su afirmación. Si un printer o captura no muestra una fecha propia, deja fecha_ejecucion_acreditada vacía: no le atribuyas la fecha de la carta. Para el contraste temporal, la fecha de la carta que remitió ese medio sin fecha se utilizará después como fecha de referencia.

REGLA TEMPORAL OBLIGATORIA: distingue siempre la fecha del documento o de su remisión de la fecha efectiva de ejecución que el documento acredita. Un documento posterior al vencimiento puede acreditar una ejecución anterior únicamente cuando contiene un registro histórico, fecha de operación, evento de sistema u otro dato objetivo que identifique esa ejecución anterior. Un printer que solo muestra el estado actual "activo", sin fecha histórica de reconexión o de estado, no acredita por sí solo que la obligación se ejecutó dentro del plazo. Si el medio carece de fecha efectiva, deja fecha_ejecucion_acreditada vacía; nunca copies allí automáticamente la fecha de la carta.

RECONEXIÓN CONDICIONADA: si el mandato ordena reconectar solamente cuando el servicio se encuentre suspendido, verifica con evidencia fechada si la condición existía en la fecha de notificación y durante el plazo. Para concluir condicion_no_configurada debe existir un histórico o consulta fechada que acredite que el servicio ya estaba activo en esa fecha relevante. Que aparezca activo en una consulta posterior no demuestra por sí solo que nunca estuvo suspendido ni que se reconectó oportunamente.

INVENTARIO DE EVIDENCIA: distingue entre un medio no presentado y un medio presentado cuyo contenido no acredita el hecho o cuya fecha no es verificable. Si un archivo, índice o página identifica expresamente un "Histórico de cortes y reconexiones", "Consulta del estado del servicio" o printer equivalente, regístralo como presentado y está prohibido afirmar que la empresa "no lo remitió". Si sus datos no muestran una fecha útil, indica con precisión que fue presentado pero no permite verificar temporalmente la reconexión.

REGLA OBLIGATORIA PARA ENTREGA DE CONTRATOS O INFORMACIÓN: una carta o correo simple no acredita por sí solo la ejecución. En entrega física exige cargo de notificación al domicilio con fecha y forma de entrega. En correo electrónico exige: (a) constancia de envío y (b) cualquiera de estos dos medios alternativos: constancia de entrega al destinatario o al buzón de destino, o constancia de recepción del usuario; todos con fecha verificable. No exijas acumular constancia de entrega y constancia de recepción del usuario: cualquiera de ellas, acompañada por la constancia de envío, acredita la entrega. Tampoco exijas confirmación de lectura o respuesta del usuario. El contrato meramente adjunto sin constancia de envío ni alguno de los dos medios alternativos no basta. Si falta una prueba, describe únicamente cuál no obra; está prohibido afirmar que el servidor rechazó, no envió o no confirmó algo, salvo que el documento lo diga literalmente. Toma como fecha de ejecución la fecha acreditada por la constancia de entrega o, en su defecto, por la constancia de recepción del usuario, y compárala con el vencimiento para determinar si fue oportuna o tardía. En esta materia está prohibido usar razonamientos sobre periodos, meses, descuentos, registro o activación, salvo que esos elementos formen parte real del mandato evaluado.

EXCEPCIÓN OBLIGATORIA — AJUSTES, ANULACIONES O DESCUENTOS EN LA FACTURACIÓN: cuando la obligación consiste en ajustar, anular o descontar un importe en la facturación (no en devolver dinero en efectivo), la ejecución se acredita con la captura de pantalla del sistema o el histórico del estado de cuenta que muestre que el ajuste coincide con el importe ordenado por el TRASU, conforme a los criterios de la materia "Facturación y cobro". En estos casos NO se exige acreditar que el usuario recibió una notificación o carta sobre el ajuste; dicha comunicación, si existe, es evidencia adicional pero no condición de cumplimiento. No confundas la obligación de ajustar la facturación con una obligación de informar al usuario.

No evalúes todavía PAS ni subsanación. Devuelve JSON conforme al esquema."""
    extraction_raw=parse_json_response(gemini_text(extraction_system,json.dumps({"esquema":extraction_schema,"caso":payload},ensure_ascii=False),json_mode=True))
    extraction=json_object(extraction_raw,("extraccion","extraction","resultado"))
    if not extraction: raise ValueError("Gemini no devolvió una extracción jurídica utilizable")
    schema={"ficha":{"expediente":"","empresa_operadora":"","tipo_acto":"","numero_acto":"","fecha_notificacion_emision":"","plazo_cumplimiento":"","fecha_vencimiento":"","obligacion_principal":"","medios_probatorios":[]},"trazabilidad":[],"evaluacion_juridica":{"checklist":[],"sustento_breve":"","tipo_incumplimiento":""},"resultado":"Cumplió|Incumplió|Inejecutable","subsanacion_voluntaria":"aplica|no aplica|no corresponde","clasificacion":"PAS|NO PAS","parrafo_final":"","datos_pendientes":[]}
    system="""Eres analista jurídico senior de OSIPTEL. Usa SOLO la evidencia y fuentes entregadas. No inventes fechas, pruebas ni conclusiones. Lo faltante es 'No identificado' o 'Pendiente de verificación'.

JERARQUÍA DOCUMENTAL OBLIGATORIA:
1. Sigue literalmente instrucciones_juridicas.txt para determinar el tipo de acto, el orden del análisis, la prueba exigible y los datos que no pueden inferirse.
2. El sistema ya consultó criterios_evaluacion_obligatorios.txt y seleccionó, desde la obligación principal, la materia y las filas aplicables. Aplica íntegramente SOLO esas filas como instrucciones vinculantes. No uses filas no seleccionadas ni conocimiento general para sustituirlas. Conserva exactamente el significado de PAS, NO PAS, PLAZO e INEJECUTABLE de cada fila.
3. Aplica como reglas vinculantes las filas seleccionadas de PAUTAS PAS.xlsx, separando cumplimiento, razonabilidad, cese y subsanación voluntaria.
4. Redacta con la estructura y lenguaje de PLANTILLAS cumplimiento.docx cuando sea una Resolución TRASU.
5. Redacta con la estructura y lenguaje de PLANTILLAS DENUNCIAS ACTUALIZADAS.docx para cartas, denuncias, SAR, SARA o SAP.
6. Antes de concluir, identifica en el checklist el nombre del archivo y la fila o sección aplicada. Si no puedes identificar la regla utilizada, no emitas evaluación final y registra el dato como pendiente.
7. Está prohibido emitir una conclusión basada solamente en conocimiento general del modelo cuando contradiga cualquiera de esos documentos.

MÉTODO Y CONTROLES JURÍDICOS OBLIGATORIOS (en este orden):
0. Determina la materia y cita en el checklist las filas concretas de criterios y pautas usadas. No uses una regla de otra materia.
1. Evalúa cada componente del mandato y, únicamente cuando el propio mandato tenga periodos, evalúa cada periodo. El cumplimiento parcial NO equivale a cumplimiento íntegro.
2. Determina la prueba exigible y el resultado únicamente con las filas seleccionadas de la materia aplicable y con los hechos acreditados.
3. Si el mandato contiene varios componentes, evalúa cada componente sin importar criterios de otra materia.
4. La subsanación voluntaria exige conjuntamente cese TOTAL de la conducta y reversión INTEGRAL de todos sus efectos antes del procedimiento. Que la materia no restrinja el servicio, por sí solo, jamás basta.
4-A. Usa obligatoriamente los campos restriccion_servicio y sustento_restriccion_servicio de las fuentes. Si restriccion_servicio=NO, está prohibido afirmar que el mandato está vinculado a una restricción del servicio o descartar el eximente por efectos irreversibles de una restricción. En ese supuesto evalúa separadamente cese total, reversión integral y oportunidad anterior al inicio del PAS. Si restriccion_servicio=SI, explica la concreta afectación al acceso o uso material del servicio.
5. Si siguen componentes pendientes, explica la ausencia de cese total o reversión integral según las pautas seleccionadas.
6. Usa exclusivamente la fecha de vencimiento calculada y no la recalcules.
7. Aplica las instrucciones, criterios y pautas PAS entregados, incluso si contradicen una inferencia general del modelo. La plantilla determina la estructura de redacción.
8. Contrasta obligación por obligación con la matriz de pruebas; no generalices el resultado de un componente a los demás.
9. Mantén separadas tres decisiones: (a) cumplimiento material dentro del plazo, (b) razonabilidad para una ejecución tardía y (c) eximente de subsanación voluntaria. No conviertas automáticamente una ejecución tardía en subsanación.
10. No traslades pruebas, frases ni conclusiones de una materia distinta de la materia seleccionada.
11. La etiqueta NO PAS de una fila solo procede si todos los hechos descritos en esa fila están acreditados. Si faltan meses, extremos o prueba objetiva, aplica la fila específica de incumplimiento/PAS.
12. Antes de redactar, construye el checklist con: mandato; plazo; prueba exigible; prueba existente; componentes acreditados; componentes pendientes; regla aplicada; conclusión. Incluye periodos acreditados o pendientes solo cuando el mandato realmente los contenga. Si existe contradicción, prevalece la evidencia documental y la regla específica.
13. El párrafo final debe indicar obligatoriamente y de forma expresa: fecha de notificación, número y tipo de días del plazo verificado, y fecha de vencimiento. Copia esos tres datos literalmente de la ficha; está prohibido recalcularlos u omitir el plazo.
14. Cada conclusión debe derivarse de hechos mencionados inmediatamente antes. No concluyas cumplimiento, incumplimiento, cese, reversión o subsanación si la matriz no identifica la prueba y los componentes que sustentan esa conclusión. Exige periodos solo cuando formen parte de la obligación.
15. El párrafo final debe seguir literalmente la redacción y estructura de frases de la plantilla aplicable (PLANTILLAS cumplimiento.docx o PLANTILLAS DENUNCIAS ACTUALIZADAS.docx). Está prohibido agregar datos que la plantilla no contempla en esa oración, como el número de la resolución, carta, SARA, SAR, SAP o resolución de primera instancia, salvo que la plantilla lo incluya expresamente en su texto.
16. Está prohibido afirmar que una obligación se ejecutó "dentro del plazo" si la extracción probatoria no contiene una fecha_ejecucion_acreditada igual o anterior a la fecha de vencimiento. Distingue fecha de carta, fecha de captura y fecha histórica de ejecución. Para una reconexión condicionada, un estado "activo" consultado después del vencimiento solo acredita el estado en esa fecha posterior; no acredita la inexistencia de suspensión en la fecha de notificación ni una reconexión oportuna, salvo que el propio histórico identifique esas fechas.
17. CONTROL DE MATERIA: aplica únicamente criterios correspondientes a la obligación principal identificada. Si la obligación es reconectar, reactivar o acreditar operatividad, está prohibido fundamentar el análisis con reglas o expresiones sobre descuentos recurrentes, meses o periodos pendientes, importes, notas de crédito, registro o activación de beneficios, ofertas o promociones. Esas expresiones solo proceden cuando el mandato principal versa realmente sobre esas materias. Antes de redactar, elimina del razonamiento cualquier criterio perteneciente a una materia distinta.
18. CONTROL DE EXISTENCIA DOCUMENTAL: no confundas ausencia de un documento con insuficiencia de su contenido. Si el expediente o la extracción probatoria identifica que se presentó un histórico, consulta, printer, acta o constancia, menciona que fue presentado. Solo concluye que no acredita el cumplimiento cuando falte en ese medio la fecha, el evento o el dato objetivo exigible. Nunca escribas "no remitió" o "no adjuntó" respecto de un medio que aparece en el inventario documental.
19. CONTROL DE ENTREGA DE CONTRATOS O INFORMACIÓN: si el mandato consiste en entregar documentos o brindar información, analiza exclusivamente el contenido ordenado, el medio utilizado y la prueba de entrega. Para entrega física exige cargo de notificación al domicilio con fecha y forma de entrega. Para correo electrónico exige constancia de envío y, adicionalmente, uno de estos medios alternativos: constancia de entrega al destinatario o constancia de recepción del usuario, con fecha verificable. No exijas ambos medios alternativos ni confirmación de lectura o respuesta. Si falta la constancia de envío o ambos medios alternativos, escribe exactamente qué no obra; no atribuyas al servidor una negativa, rechazo o falta de confirmación que no esté literalmente documentada. Usa como fecha de ejecución la fecha acreditada por la constancia de entrega o, en su defecto, por la constancia de recepción, y compárala con el vencimiento. Está prohibido trasladar a estos casos frases sobre periodos, meses pendientes, descuentos, registro o activación.
20. CITA OBLIGATORIA DE MEDIOS PROBATORIOS: por cada componente o periodo que menciones como acreditado, parcialmente acreditado o pendiente, cita el tipo exacto de documento tal como figura en extraccion_probatoria.medios_probatorios[].documento (por ejemplo "nota de crédito", "printer", "constancia", "carta", "histórico"), junto con su fecha (fecha_documento o, si corresponde a la ejecución, fecha_ejecucion_acreditada). Está prohibido sustituir ese nombre específico por términos genéricos como "documentos", "la documentación remitida" o "los medios probatorios" cuando extraccion_probatoria contenga el nombre concreto del documento. Cita siempre la denominación que consta en el contenido del documento; está PROHIBIDO citar nombres de archivos informáticos o sus extensiones (.pdf, .tif, .jpg, etc.), pues el nombre del archivo es irrelevante y no identifica al documento.
21. PROHIBICIÓN DE INVENTAR CAUSAS O CALIFICACIONES: al describir la razón de un componente pendiente o incumplido, usa exclusivamente los hechos y el lenguaje que constan literalmente en extraccion_probatoria (campo hecho_acreditado o cita). Está prohibido acuñar o introducir calificaciones técnicas o jurídicas ausentes del documento fuente (por ejemplo "restricción sistémica", "limitación técnica", "caso fortuito", "fuerza mayor") aunque parezcan un resumen razonable. Si la fuente dice que algo está "en gestión" o "programado en el sistema", redacta con esas mismas palabras o una paráfrasis mínima y fiel, nunca con una etiqueta nueva.

Devuelve JSON válido conforme al esquema y un párrafo final completo, cronológico, con obligación, pruebas por cada componente citando el documento exacto y su fecha, contraste, conclusión y análisis separado de subsanación. Solo usa periodos cuando el mandato los contenga."""
    user=json.dumps({"esquema":schema,"caso":{k:v for k,v in payload.items() if k!="documentos"},"extraccion_probatoria":extraction,"fuentes":sources},ensure_ascii=False)
    result_raw=parse_json_response(gemini_text(system,user,json_mode=True))
    result=json_object(result_raw,("evaluacion","evaluation","resultado"))
    if not result: raise ValueError("Gemini no devolvió una evaluación jurídica utilizable")
    if not isinstance(result.get("ficha"),dict): result["ficha"]={}
    if not isinstance(result.get("evaluacion_juridica"),dict): result["evaluacion_juridica"]={}
    checklist=result["evaluacion_juridica"].get("checklist")
    if not isinstance(checklist,list): checklist=[]
    selection_note=("criterios_evaluacion_obligatorios.txt — materia: "
                    f"{sources['materia_criterios_seleccionada']}; filas consultadas: "
                    f"{sources['filas_criterios_seleccionadas']}.")
    if not any("criterios_evaluacion_obligatorios.txt" in str(item) for item in checklist):
        checklist.insert(0,selection_note)
    result["evaluacion_juridica"]["checklist"]=checklist
    trace=result.get("trazabilidad")
    if not isinstance(trace,list): trace=[]
    trace.append(selection_note)
    result["trazabilidad"]=trace
    # Verified spreadsheet values are authoritative; the model cannot recalculate them.
    result.setdefault("ficha",{})["expediente"]=payload.get("expediente_detectado","No identificado")
    result["ficha"]["tipo_acto"]=payload.get("tipo_acto","No identificado")
    result["ficha"]["fecha_notificacion_emision"]=payload.get("fecha_verificada","No identificado")
    result["ficha"]["plazo_cumplimiento"]=payload.get("plazo_verificado","Pendiente de verificación")
    result["ficha"]["fecha_vencimiento"]=payload.get("fecha_vencimiento","Pendiente de verificación")
    verified_operator=str(payload.get("empresa_operadora_verificada") or "").strip()
    if verified_operator:
        result["ficha"]["empresa_operadora"]=verified_operator
    verified_principal=payload.get("obligacion_y_plazo_principales_verificados") or {}
    resolutive_obligation=extraction.get("obligacion_extraida_parte_resolutiva") or {}
    if isinstance(verified_principal,dict) and str(verified_principal.get("obligacion_principal","")).strip():
        result["ficha"]["obligacion_principal"]=str(verified_principal["obligacion_principal"]).strip()
    elif isinstance(resolutive_obligation,dict) and str(resolutive_obligation.get("texto","")).strip():
        result["ficha"]["obligacion_principal"]=str(resolutive_obligation["texto"]).strip()
    elif payload.get("tipo_acto")=="Resolución TRASU":
        raise ValueError("Se leyó HA RESUELTO, pero no fue posible distinguir con seguridad la obligación principal de los deberes accesorios")
    # Deterministic legal guard: pending/programmed work is not full execution.
    statuses={str(x.get("estado","")).lower() for x in (extraction.get("medios_probatorios") or []) if isinstance(x,dict)}
    matrix={str(x.get("estado","")).lower() for x in (extraction.get("matriz_cumplimiento") or []) if isinstance(x,dict)}
    # A single unaccredited piece of evidence (e.g. an incidental notification letter) must not
    # override compliance by itself; only the per-obligation matrix (which the main evaluation
    # step builds while reading criterios_evaluacion_obligatorios.txt) decides that. "programado"/
    # "en_curso" are the exception: they mean the action was never finished, regardless of which
    # document reports it.
    incomplete=bool(statuses & {"programado","en_curso"} or matrix & {"parcial","no_acreditado"})
    if incomplete:
        obligation_text=str(result.get("ficha",{}).get("obligacion_principal") or "")
        information_delivery=is_information_delivery_obligation(obligation_text)
        result["resultado"]="Incumplió"
        result["subsanacion_voluntaria"]="no aplica"
        result["clasificacion"]="PAS"
        incomplete_note=(
            "Regla obligatoria de entrega de información: uno o más componentes del contenido, la entrega o la "
            "recepción no están acreditados; no hubo ejecución íntegra."
            if information_delivery else
            "Regla obligatoria: uno o más componentes de la obligación están programados, en curso o no "
            "acreditados; no hubo ejecución íntegra.")
        result.setdefault("evaluacion_juridica",{}).setdefault("checklist",[]).append(incomplete_note)
        locked={
            "ficha":result.get("ficha",{}),"extraccion_probatoria":extraction,
            "resultado_obligatorio":"Incumplió","subsanacion_obligatoria":"no aplica",
            "clasificacion_obligatoria":"PAS","fuentes_aplicables":sources,
            "tipo_obligacion":"entrega_de_informacion" if information_delivery else "otra",
        }
        rewrite_system="""Redacta un único párrafo jurídico conforme a la plantilla TRASU. Las conclusiones indicadas como obligatorias están bloqueadas y no puedes modificarlas. Debes indicar literalmente la fecha de notificación, el plazo de cumplimiento (número y tipo de días) y la fecha de vencimiento que aparecen en la ficha; no puedes recalcularlos ni omitirlos. No afirmes que una programación, solicitud, carta o gestión en curso acredita ejecución. CITA OBLIGATORIA DE MEDIOS PROBATORIOS: por cada componente o periodo (acreditado, parcial o pendiente), cita el tipo exacto de documento tal como figura en extraccion_probatoria.medios_probatorios[].documento (por ejemplo "nota de crédito", "printer", "constancia", "carta", "histórico") junto con su fecha (fecha_documento o fecha_ejecucion_acreditada, la que corresponda); está prohibido reemplazar ese nombre específico por términos genéricos como "documentos" o "la documentación remitida" cuando extraccion_probatoria contenga el nombre concreto. Cita siempre la denominación que consta en el contenido del documento; está PROHIBIDO citar nombres de archivos informáticos o sus extensiones (.pdf, .tif, .jpg, etc.), pues el nombre del archivo no identifica al documento. PROHIBICIÓN DE INVENTAR CAUSAS: al explicar por qué un componente sigue pendiente, usa exclusivamente los hechos y el lenguaje literal de extraccion_probatoria (hecho_acreditado o cita); está prohibido acuñar calificaciones técnicas o jurídicas ausentes de la fuente (p. ej. "restricción sistémica", "limitación técnica", "caso fortuito", "fuerza mayor") aunque parezcan un resumen razonable — si la fuente dice "en gestión" o "programado", usa esas mismas palabras o una paráfrasis mínima y fiel. Si tipo_obligacion=entrega_de_informacion, analiza únicamente el contenido ordenado y su entrega: la notificación física se acredita con cargo al domicilio; el correo electrónico requiere constancia de envío y, alternativamente, constancia de entrega al destinatario o constancia de recepción del usuario. No exijas acumular ambos medios alternativos ni confirmación de lectura o respuesta. Usa como fecha de ejecución la fecha de entrega o, en su defecto, la fecha de recepción acreditada, y determina si fue oportuna o tardía. Cuando falte una prueba escribe solo cuál no obra; no inventes respuestas del servidor. En esos casos está prohibido mencionar periodos, meses pendientes, descuentos, registro o activación. Para las demás obligaciones explica de forma neutral qué componente material quedó pendiente, sin importar frases de otra materia. No apliques el eximente por el solo hecho de que no exista restricción del servicio. Devuelve JSON {parrafo_final:string}."""
        rewritten=parse_json_response(gemini_text(rewrite_system,json.dumps(locked,ensure_ascii=False),json_mode=True)) or {}
        if not isinstance(rewritten,dict): rewritten={}
        result["parrafo_final"]=rewritten.get("parrafo_final",result.get("parrafo_final",""))
    # El dictamen nunca puede quedar vacío ni fuera del vocabulario cerrado.
    normalize_verdict_fields(result,extraction,sources)
    # A response without a paragraph is never a completed evaluation.
    if not str(result.get("parrafo_final","")).strip():
        locked={"ficha":result.get("ficha",{}),"extraccion_probatoria":extraction,
                "resultado":result.get("resultado"),"subsanacion_voluntaria":result.get("subsanacion_voluntaria"),
                "clasificacion":result.get("clasificacion"),"fuentes_aplicables":sources}
        fallback_system="""Redacta el párrafo jurídico final conforme a la plantilla aplicable. Relaciona en orden: acto y notificación; plazo y vencimiento de la ficha; mandato; alegaciones; pruebas objetivas citando el tipo exacto de documento (según extraccion_probatoria.medios_probatorios[].documento, p. ej. "nota de crédito", "printer", "constancia") y su fecha, sin usar términos genéricos como "documentos" cuando el nombre concreto conste en extraccion_probatoria, y sin citar jamás nombres de archivos informáticos ni extensiones (.pdf, .tif, etc.), pues el documento se identifica por la denominación de su contenido; contraste por cada obligación o periodo; conclusión; y subsanación. No cambies el resultado, la clasificación ni la subsanación ya determinadas. No inventes causas, calificaciones técnicas o jurídicas (p. ej. "restricción sistémica", "caso fortuito") que no consten literalmente en extraccion_probatoria; usa exactamente el lenguaje de la fuente. Devuelve JSON {parrafo_final:string}."""
        try:
            fallback=parse_json_response(gemini_text(fallback_system,json.dumps(locked,ensure_ascii=False),json_mode=True)) or {}
            if isinstance(fallback,dict): result["parrafo_final"]=str(fallback.get("parrafo_final","")).strip()
        except Exception:
            result["parrafo_final"]=""
    if not str(result.get("parrafo_final","")).strip():
        result["parrafo_final"]=deterministic_paragraph(result,extraction)
    paragraph=str(result.get("parrafo_final","")).strip()
    if payload.get("tipo_acto")=="Resolución TRASU":
        # The applicable template starts with "La Resolución TRASU" and does
        # not reproduce the identifying number in the final paragraph.
        notification=str(payload.get("fecha_verificada")).strip()
        paragraph=re.sub(r"^La\s+Resoluci[oó]n(?:\s+TRASU)?(?:\s+N(?:ro\.?|[.°º])?\s*[0-9][A-Z0-9./-]*)?\s+fue\s+(?:emitida|notificada)(?:\s+en\s+fecha\s+no\s+identificada|\s+el\s+[^,.;]+)",
                         f"La Resolución TRASU fue notificada el {notification}",paragraph,count=1,flags=re.I)
        paragraph=re.sub(r"^La\s+Resoluci[oó]n(?:\s+TRASU)?\s+N(?:ro\.?|[.°º])?\s*[0-9][A-Z0-9./-]*",
                         "La Resolución TRASU",paragraph,count=1,flags=re.I)
    # The templates never cite the act's identifying number; strip it if the model added it anyway.
    numero_acto=str(result.get("ficha",{}).get("numero_acto","")).strip()
    if numero_acto and numero_acto.lower() not in {"no identificado","pendiente de verificación","pendiente de verificacion"}:
        paragraph=re.sub(r"\s*N\.?[°º]\s*"+re.escape(numero_acto),"",paragraph)
    # Avoid an accidental exact duplication of the generated paragraph.
    half=len(paragraph)//2
    if len(paragraph)%2==0 and paragraph[:half]==paragraph[half:]:
        paragraph=paragraph[:half].strip()
    paragraph=enforce_conditional_reconnection_timeline(
        result,extraction,paragraph,sources,payload.get("documentos") if isinstance(payload.get("documentos"),dict) else {})
    paragraph=enforce_service_restriction_rule(result,paragraph,sources,extraction)
    paragraph=enforce_information_delivery_language(
        result,paragraph,payload.get("documentos") if isinstance(payload.get("documentos"),dict) else {})
    paragraph=enforce_evidence_citation(extraction,paragraph)
    derive_execution_fields(result,extraction)
    result["parrafo_final"]=normalize_legal_paragraph(paragraph,result.get("ficha",{}),result.get("resultado",""))
    # Estos campos personales no son necesarios para la evaluación ni deben
    # conservarse en el estado, historial o exportación de CumpleTRASU.
    if isinstance(result.get("ficha"),dict):
        result["ficha"].pop("usuario_abonado",None)
        result["ficha"].pop("servicio",None)
    return result

def regenerate_paragraph(result: dict[str,Any]) -> str:
    context={"ficha":result.get("ficha",{}),"evaluacion_juridica":result.get("evaluacion_juridica",{}),"resultado":result.get("resultado"),"subsanacion_voluntaria":result.get("subsanacion_voluntaria"),"clasificacion":result.get("clasificacion"),"datos_pendientes":result.get("datos_pendientes",[]),"instrucciones":INSTRUCCIONES.read_text("utf-8",errors="ignore")[:40000]}
    response=gemini_text("Regenera únicamente el párrafo jurídico final con los datos aportados. Usa todas las fechas exclusivamente en formato dd/mm/aaaa. Si es TRASU, inicia con 'La Resolución TRASU fue notificada el dd/mm/aaaa' y no incluyas el número de resolución ni sustituyas la notificación por la emisión. Respeta el campo evaluacion_juridica.restriccion_servicio: cuando sea 'No', está prohibido afirmar que el mandato restringe el servicio o que sus efectos son irreversibles por ese motivo. Si la obligación es entregar contratos o información por correo, exige constancia de envío y, como prueba alternativa de llegada, constancia de entrega al destinatario o constancia de recepción del usuario; no exijas ambas ni confirmación de lectura o respuesta. Usa la fecha acreditada de entrega o recepción para determinar si la ejecución fue oportuna o tardía. No inventes respuestas del servidor ni uses frases sobre periodos, descuentos, registro o activación. No inventes ni completes faltantes. Devuelve JSON {parrafo_final:string}.",json.dumps(context,ensure_ascii=False),json_mode=True)
    paragraph=parse_json_response(response)["parrafo_final"]
    # El párrafo regenerado también debe respetar la regla de restricción del
    # servicio; antes este camino omitía el control y podía reintroducir una
    # restricción falsa en materias informativas.
    stored=_fold_legal_text(result.get("evaluacion_juridica",{}).get("restriccion_servicio"))
    pseudo_sources={"restriccion_servicio":"SI" if stored.startswith("si") else "NO",
                    "sustento_restriccion_servicio":""}
    paragraph=enforce_service_restriction_rule(result,paragraph,pseudo_sources)
    return normalize_legal_paragraph(paragraph,result.get("ficha",{}),result.get("resultado",""))

def to_row(result: dict, documents: list[str]) -> dict:
    f=result.get("ficha",{}); e=result.get("evaluacion_juridica",{})
    return dict(zip(COLUMNAS,[f.get("expediente",""),f.get("empresa_operadora",""),f.get("tipo_acto",""),f.get("numero_acto",""),f.get("fecha_notificacion_emision",""),f.get("fecha_vencimiento",""),f.get("obligacion_principal",""),f.get("estado_ejecucion",""),f.get("fecha_ejecucion",""),result.get("resultado",""),e.get("tipo_incumplimiento",""),result.get("subsanacion_voluntaria",""),result.get("clasificacion",""),e.get("sustento_breve",""),result.get("parrafo_final",""),"; ".join(map(str,result.get("datos_pendientes",[]))),"; ".join(documents),datetime.now().strftime("%Y-%m-%d %H:%M")]))

def excel_bytes(row: dict | None=None) -> bytes:
    existing=pd.read_excel(HISTORIAL,dtype=str) if HISTORIAL.exists() else pd.DataFrame(columns=COLUMNAS)
    existing=existing.drop(columns=["Usuario o abonado","Servicio"],errors="ignore").reindex(columns=COLUMNAS)
    if row is not None: existing=pd.concat([existing,pd.DataFrame([row])],ignore_index=True)
    b=io.BytesIO()
    with pd.ExcelWriter(b,engine="openpyxl") as w: existing.to_excel(w,index=False,sheet_name="Evaluaciones")
    return b.getvalue()

for k,v in {"result":None,"docs":[],"texts":{},"notice":None,"due":None,
            "analysis_error":None,"analysis_status":None,"upload_signature":None,
            "debug_resolutivo":None,"uploader_version":0}.items():
    st.session_state.setdefault(k,v)

left,center,right=st.columns([1.05,2.25,1.05],gap="large")
with left:
    st.markdown('<div class="light">Expediente particular</div>',unsafe_allow_html=True)
    uploads=st.file_uploader(
        "Cargar expediente",accept_multiple_files=True,
        type=["pdf","docx","xlsx","xls","csv","txt","png","jpg","jpeg","tif","tiff","zip","rar","7z"],
        key=f"expediente_upload_{st.session_state.uploader_version}")
    if uploads:
        signature="|".join(f"{u.name}:{hashlib.sha256(u.getvalue()).hexdigest()}" for u in uploads)
        if signature!=st.session_state.upload_signature:
            st.session_state.result=None; st.session_state.analysis_error=None
            st.session_state.analysis_status="Archivos nuevos listos para analizar"
            st.session_state.upload_signature=signature
        st.session_state.docs=[u.name for u in uploads]
        for u in uploads:
            kind=classify(u.name,""); st.caption(f"✓ {u.name} · {kind}")
    ok_sources,missing=fuente_status()
    st.divider(); st.markdown('<div class="light">Fuentes permanentes</div>',unsafe_allow_html=True)
    st.success("Fuentes + instrucciones jurídicas y criterios obligatorios disponibles") if ok_sources else st.error(f"Faltan {len(missing)} fuentes o instrucciones")
    if missing:
        with st.expander("Ver faltantes"): st.write("\n".join(f"• {x}" for x in missing))
    analyze=st.button("✦ Analizar expediente con IA",type="primary",use_container_width=True,disabled=not uploads)

if analyze:
    st.session_state.result=None; st.session_state.analysis_error=None
    st.session_state.analysis_status="Análisis iniciado"
    if not ok_sources: st.error("La herramienta no puede emitir evaluación final porque no tiene acceso a las fuentes permanentes")
    elif not get_api_key(): st.error("Configure GEMINI_API_KEY en .env o en los secretos de Streamlit.")
    else:
        with st.spinner("Procesando y contrastando el expediente..."):
            case_dir=Path(tempfile.mkdtemp(prefix="caso_",dir=TEMP))
            try:
                paths=[]
                for u in uploads:
                    p=case_dir/safe_name(u.name); p.write_bytes(u.getbuffer()); paths.append(p)
                paths+=extract_archives(case_dir)
                unique_paths=[]; seen_hashes=set()
                for p in paths:
                    if p.suffix.lower() in {".zip",".rar",".7z"}: continue
                    try: digest=hashlib.sha256(p.read_bytes()).hexdigest()
                    except Exception: digest=str(p.resolve())
                    if digest not in seen_hashes:
                        seen_hashes.add(digest); unique_paths.append(p)
                paths=unique_paths
                texts={p.name:read_file(p) for p in paths if p.suffix.lower() not in {".zip",".rar",".7z"}}
                # A document whose OCR/reading failed must never be treated as usable
                # content — otherwise the AI can latch onto an unrelated document and
                # fabricate an obligation instead of honestly reporting "not found".
                failed_docs=[n for n,t in texts.items() if str(t).startswith("[OCR no disponible") or str(t).startswith("[Error al leer")]
                failed_details=" | ".join(str(texts[n])[:1200] for n in failed_docs)
                for n in failed_docs: texts[n]=""
                if not texts or not any(str(x).strip() for x in texts.values()):
                    raise ValueError("No se pudo extraer texto de los archivos del expediente"+(f" (falló la lectura de: {', '.join(failed_docs)})" if failed_docs else ""))
                st.session_state.analysis_status="Documentos leídos; verificando expediente y plazo"
                combined="\n\n".join(f"### {n}\n{t}" for n,t in texts.items())
                document_types=[classify(n,t) for n,t in texts.items()]
                priority=["Resolución TRASU","SARA","SAR","SAP","Denuncia","Carta","Resolución de primera instancia"]
                unmistakable_trasu=(any("TRASU" in str(n).upper() for n in texts) or
                                    bool(re.search(r"\bTRASU\b",combined,re.I)))
                tipo="Resolución TRASU" if unmistakable_trasu else next((candidate for candidate in priority if candidate in document_types),"No identificado")
                resolutive=extract_resolutive_part(texts) if tipo=="Resolución TRASU" else None
                st.session_state["debug_resolutivo"]=resolutive
                if tipo=="Resolución TRASU" and not resolutive:
                    raise ValueError("No se identificó la sección HA RESUELTO completa; no es posible establecer la obligación principal sin esa sección"+(f". Detalle de lectura: {failed_details}" if failed_details else ""))
                principal_data=extract_trasu_mandate_and_term(resolutive) if resolutive else {}
                searchable="\n".join(u.name for u in uploads)+"\n"+combined
                expediente=parse_trasu_name(searchable) or identify_exact_expediente(searchable) or "No identificado"
                notice=None; due=None; term=None; operator=None
                if tipo=="Resolución TRASU":
                    case_record=exact_case_record(expediente) if expediente!="No identificado" else None
                    notice=(case_record or {}).get("fecha_notificacion")
                    operator=(case_record or {}).get("empresa_operadora")
                    if not notice:
                        st.session_state.analysis_error="No se encontró coincidencia exacta del expediente en notificaciones 2026. Expediente detectado: "+expediente
                    else:
                        deadline=calculate_due(notice,str(principal_data.get("plazo_principal_textual") or ""))
                        if deadline: due,term=deadline
                        if not deadline: st.session_state.analysis_error="No se pudo determinar el plazo o calcular el vencimiento con CONTADOR DE PLAZOS - TRASU 2026.xlsx"
                if tipo!="Resolución TRASU" or (notice and due):
                    st.session_state.analysis_status="Aplicando instrucciones, criterios, pautas y plantillas"
                    payload={"tipo_acto":tipo,"expediente_detectado":expediente,"empresa_operadora_verificada":operator or "","fecha_verificada":notice or "No identificado","plazo_verificado":term or "Pendiente de verificación","fecha_vencimiento":due or "Pendiente de verificación","parte_resolutiva_trasu":resolutive or "No corresponde","obligacion_y_plazo_principales_verificados":principal_data,"documentos":texts}
                    st.session_state.result=ai_evaluate(payload); st.session_state.texts=texts
                    st.session_state.analysis_status="Evaluación completada"
            except Exception as e:
                st.session_state.analysis_error=f"No se pudo completar el análisis: {type(e).__name__}: {e}"
                st.session_state.analysis_status=""
            finally: shutil.rmtree(case_dir,ignore_errors=True)

if st.session_state.analysis_error:
    st.error(st.session_state.analysis_error)
if st.session_state.debug_resolutivo:
    with st.expander("Ver texto de la parte resolutiva detectada (para diagnóstico)"):
        st.text(st.session_state.debug_resolutivo)
elif st.session_state.analysis_status:
    st.info(st.session_state.analysis_status)

r=st.session_state.result
with center:
    st.markdown('<div class="light">Ficha editable del expediente</div>',unsafe_allow_html=True)
    f=(r or {}).get("ficha",{})
    a,b=st.columns(2)
    expediente=a.text_input("Expediente",f.get("expediente","No identificado")); empresa=b.text_input("Empresa operadora",f.get("empresa_operadora","No identificado"))
    tipo=a.text_input("Tipo de acto",f.get("tipo_acto","No identificado")); numero=b.text_input("Número de resolución o carta",f.get("numero_acto","No identificado"))
    notif=a.text_input("Fecha de notificación o emisión",f.get("fecha_notificacion_emision","No identificado")); plazo=b.text_input("Plazo de cumplimiento",f.get("plazo_cumplimiento","Pendiente de verificación"))
    vence=st.text_input("Fecha máxima de vencimiento",f.get("fecha_vencimiento","Pendiente de verificación"))
    execution_options=["Ejecutó","No ejecutó","No aplica","Pendiente de verificación"]
    current_execution=f.get("estado_ejecucion","Pendiente de verificación")
    if current_execution not in execution_options: current_execution="Pendiente de verificación"
    estado_ejecucion=a.selectbox("Estado de ejecución",execution_options,index=execution_options.index(current_execution))
    fecha_ejecucion=b.text_input("Fecha de ejecución acreditada",f.get("fecha_ejecucion","Pendiente de verificación"))
    obligacion=st.text_area("Obligación principal",f.get("obligacion_principal","No identificado"),height=90)
    st.markdown("**Medios probatorios**")
    medios=f.get("medios_probatorios",[]) if r else []
    st.write("\n".join(f"• {x}" for x in medios) if medios else "No identificado")
    with st.expander("Checklist jurídico",expanded=True):
        checks=(r or {}).get("evaluacion_juridica",{}).get("checklist",[])
        st.write("\n".join(f"• {x}" for x in checks) if checks else "Pendiente de evaluación")
    with st.expander("Trazabilidad"):
        trace=(r or {}).get("trazabilidad",[])
        st.json(trace) if trace else st.write("Pendiente de evaluación")

with right:
    st.markdown('<div class="light">Dictamen</div>',unsafe_allow_html=True)
    resultado=(r or {}).get("resultado","Pendiente")
    css="ok" if resultado=="Cumplió" else "bad" if resultado=="Incumplió" else "wait"
    st.markdown(f'<div class="traffic {css}">{resultado}</div>',unsafe_allow_html=True)
    st.metric("Subsanación voluntaria",(r or {}).get("subsanacion_voluntaria","Pendiente"))
    st.metric("Clasificación",(r or {}).get("clasificacion","Pendiente"))
    st.markdown("**Alertas**")
    pendientes=(r or {}).get("datos_pendientes",[])
    if pendientes:
        for x in pendientes: st.warning(str(x))
    else: st.caption("Sin alertas disponibles")

st.divider(); st.markdown('<div class="light">Párrafo final</div>',unsafe_allow_html=True)
paragraph=st.text_area("Resultado listo para copiar",(r or {}).get("parrafo_final","La evaluación aparecerá aquí cuando se complete el análisis."),height=150,label_visibility="collapsed")
components.html(f"""<button onclick='navigator.clipboard.writeText(document.getElementById("p").textContent)' style='background:#005f9e;color:#fff;border:1px solid #15b8d4;border-radius:9px;padding:8px 15px;font-weight:700;cursor:pointer'>Copiar párrafo</button><span id='p' style='display:none'>{paragraph.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')}</span>""",height=45)
c1,c2,c3,c4=st.columns(4)
with c1:
    if st.button("Guardar evaluación",use_container_width=True,disabled=not bool(r)):
        f.update({"expediente":expediente,"empresa_operadora":empresa,"tipo_acto":tipo,"numero_acto":numero,"fecha_notificacion_emision":notif,"fecha_vencimiento":vence,"estado_ejecucion":estado_ejecucion,"fecha_ejecucion":fecha_ejecucion,"obligacion_principal":obligacion}); r["parrafo_final"]=paragraph
        row=to_row(r,st.session_state.docs); data=excel_bytes(row); HISTORIAL.write_bytes(data); st.success("Evaluación guardada")
with c2: st.download_button("Exportar a Excel",excel_bytes(to_row(r,st.session_state.docs)) if r else excel_bytes(),"CumpleTRASU_evaluacion.xlsx","application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",use_container_width=True)
with c3:
    if st.button("Regenerar párrafo",use_container_width=True,disabled=not bool(r)):
        f.update({"expediente":expediente,"empresa_operadora":empresa,"tipo_acto":tipo,"numero_acto":numero,"fecha_notificacion_emision":notif,"fecha_vencimiento":vence,"estado_ejecucion":estado_ejecucion,"fecha_ejecucion":fecha_ejecucion,"obligacion_principal":obligacion,"medios_probatorios":medios})
        try:
            r["parrafo_final"]=regenerate_paragraph(r); st.session_state.result=r; st.rerun()
        except Exception as e: st.error(f"No se pudo regenerar el párrafo: {e}")
with c4:
    if st.button("Limpiar expediente",use_container_width=True):
        st.session_state.uploader_version+=1
        st.session_state.result=None
        st.session_state.docs=[]
        st.session_state.texts={}
        st.session_state.notice=None
        st.session_state.due=None
        st.session_state.analysis_error=None
        st.session_state.analysis_status=None
        st.session_state.upload_signature=None
        st.session_state.debug_resolutivo=None
        st.rerun()

st.divider(); st.markdown("### Casos evaluados")
if HISTORIAL.exists(): st.dataframe(pd.read_excel(HISTORIAL,dtype=str).drop(columns=["Usuario o abonado","Servicio"],errors="ignore"),use_container_width=True,hide_index=True)
else: st.caption("Aún no hay evaluaciones guardadas.")
st.caption("CumpleTRASU asiste el análisis jurídico; la revisión profesional y la integridad de las fuentes siguen siendo obligatorias.")
